"""
O Babilaca (IA) — Assistente Inteligente para Licitações e Contratações Públicas
Sistema integrado com IA, APIs públicas e geração de documentos (Lei 14.133/2021)
"""

import streamlit as st
import requests
import pandas as pd
import json
import os
import re
import io
import hashlib
import base64
from datetime import datetime, timedelta
from fpdf import FPDF
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from dotenv import load_dotenv
    load_dotenv(os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), ".env"))
except ImportError:
    pass  # em produção (Streamlit Cloud), usa variáveis de ambiente direto

try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# ============================================================
# CONFIGURAÇÃO DA PÁGINA
# ============================================================
st.set_page_config(
    page_title="O Babilaca (IA)",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================
# CSS CUSTOMIZADO
# ============================================================
st.markdown("""
<style>
    body, .main { background-color: #001a4d; color: #ffffff; }
    .stApp { background-color: #001a4d; }

    /* SIDEBAR MODERNA */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0a0a0a 0%, #111111 50%, #0a0a0a 100%) !important;
        border-right: 3px solid #d4af37 !important;
        box-shadow: 4px 0 15px rgba(0, 0, 0, 0.5);
    }
    [data-testid="stSidebar"] [data-testid="stVerticalBlock"] { background: transparent !important; }
    [data-testid="stSidebarNav"] { display: none !important; }

    /* Título da sidebar */
    [data-testid="stSidebar"] .stMarkdown h2 {
        color: #d4af37 !important;
        font-family: 'Arial Black', sans-serif;
        font-size: 22px;
        text-align: center;
        letter-spacing: 2px;
        text-shadow: 1px 1px 3px rgba(0, 0, 0, 0.8);
        border-bottom: 2px solid #d4af37;
        padding-bottom: 0.75rem;
        margin-bottom: 1.5rem;
    }

    /* Links de navegação na sidebar */
    [data-testid="stSidebar"] a[data-testid="stPageLink-NavLink"] {
        background: linear-gradient(135deg, #1a1a1a 0%, #252525 100%) !important;
        color: #ffffff !important;
        border: 1px solid #333333 !important;
        border-radius: 10px !important;
        margin: 0.4rem 0 !important;
        padding: 0.85rem 1.2rem !important;
        font-weight: 600 !important;
        font-size: 15px !important;
        transition: all 0.3s ease !important;
        text-decoration: none !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
    }
    [data-testid="stSidebar"] a[data-testid="stPageLink-NavLink"] span {
        color: #ffffff !important;
    }
    [data-testid="stSidebar"] a[data-testid="stPageLink-NavLink"]:hover {
        background: linear-gradient(135deg, #252525 0%, #353535 100%) !important;
        border: 1px solid #d4af37 !important;
        transform: translateX(5px);
        box-shadow: 0 4px 15px rgba(212, 175, 55, 0.25) !important;
    }
    [data-testid="stSidebar"] a[data-testid="stPageLink-NavLink"]:hover span {
        color: #d4af37 !important;
    }
    /* Link ativo / página atual */
    [data-testid="stSidebar"] a[data-testid="stPageLink-NavLink"][aria-current="page"] {
        background: linear-gradient(135deg, #d4af37 0%, #c5a028 100%) !important;
        color: #0a0a0a !important;
        border: 1px solid #d4af37 !important;
        font-weight: bold !important;
        box-shadow: 0 4px 15px rgba(212, 175, 55, 0.4) !important;
    }
    [data-testid="stSidebar"] a[data-testid="stPageLink-NavLink"][aria-current="page"] span {
        color: #0a0a0a !important;
    }
    [data-testid="stSidebar"] label, [data-testid="stSidebar"] .stText, [data-testid="stSidebar"] p { color: #ffffff !important; }
    [data-testid="stSidebar"] hr { border-color: #333333 !important; margin: 1rem 0 !important; }
    .sidebar-footer { color: #666666; font-size: 11px; text-align: center; padding: 1rem 0; border-top: 1px solid #333333; margin-top: 2rem; }

    .babilaca-header {
        background: linear-gradient(135deg, #001a4d 0%, #0033cc 100%);
        padding: 1.2rem 2rem; border-radius: 12px;
        margin-bottom: 1.5rem; text-align: center;
        box-shadow: 0 6px 20px rgba(0,0,0,0.4);
    }
    .babilaca-header h1 { color: #d4af37; margin: 0; font-size: 2rem; }
    .babilaca-header p { color: #cbd5e1; margin: 0.3rem 0 0; font-size: 0.95rem; }

    .disclaimer-box {
        background: rgba(212,175,55,0.12); border: 1px solid #d4af37;
        border-radius: 8px; padding: 0.6rem 1rem; margin-bottom: 1rem;
        color: #d4af37; font-size: 0.85rem; text-align: center;
    }
    .doc-card {
        background: linear-gradient(135deg, #0a1628 0%, #132244 100%);
        border: 1px solid #1e3a5f; border-radius: 10px; padding: 1.2rem;
        margin-bottom: 1rem; box-shadow: 0 4px 12px rgba(0,0,0,0.3);
    }
    .doc-card h4 { color: #d4af37; margin-top: 0; }
    .doc-card p  { color: #cbd5e1; font-size: 0.9rem; }

    .alert-card {
        border-radius: 10px; padding: 1rem; margin: 0.5rem 0;
        border-left: 4px solid; color: #ffffff;
    }
    .alert-high   { background: rgba(220,38,38,0.15); border-color: #dc2626; }
    .alert-medium { background: rgba(245,158,11,0.15); border-color: #f59e0b; }
    .alert-low    { background: rgba(34,197,94,0.15); border-color: #22c55e; }
    .alert-info   { background: rgba(59,130,246,0.15); border-color: #3b82f6; }

    .stat-card {
        background: linear-gradient(135deg, #0a1628 0%, #132244 100%);
        border: 1px solid #1e3a5f; border-radius: 10px; padding: 1rem;
        text-align: center;
    }
    .stat-card .valor { color: #d4af37; font-size: 1.8rem; font-weight: bold; }
    .stat-card .label { color: #94a3b8; font-size: 0.8rem; }

    div[data-testid="stChatMessage"] { background: rgba(10,22,40,0.6) !important; border-radius: 10px; }

    /* Auto-scroll chat to bottom */
    .stChatFloatingInputContainer { position: sticky; bottom: 0; z-index: 100; }
    section[data-testid="stChatMessageContainer"] {
        max-height: 65vh;
        overflow-y: auto;
        display: flex;
        flex-direction: column;
    }

    /* Traduzir textos do file_uploader para PT-BR */
    [data-testid="stFileUploaderDropzone"] span:has(+ small) {
        visibility: hidden;
        position: relative;
    }
    [data-testid="stFileUploaderDropzone"] span:has(+ small)::after {
        content: "Arraste e solte o arquivo aqui";
        visibility: visible;
        position: absolute;
        left: 0;
        white-space: nowrap;
    }
    [data-testid="stFileUploaderDropzone"] small {
        visibility: hidden;
        position: relative;
    }
    [data-testid="stFileUploaderDropzone"] small::after {
        content: "Limite de 200MB por arquivo";
        visibility: visible;
        position: absolute;
        left: 0;
        white-space: nowrap;
    }
    [data-testid="stFileUploaderDropzoneInstructions"] button[data-testid="baseButton-minimal"] {
        visibility: hidden;
        position: relative;
    }
    [data-testid="stFileUploaderDropzoneInstructions"] button[data-testid="baseButton-minimal"]::after {
        content: "Procurar arquivos";
        visibility: visible;
        position: absolute;
        left: 50%;
        transform: translateX(-50%);
        white-space: nowrap;
        color: #d4af37;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================
# CONSTANTES
# ============================================================
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
MEMORIA_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "babilaca_memoria.json")
REQ_COUNTER_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "babilaca_req_count.json")

def _ler_req_data() -> dict:
    """Lê dados persistentes de requisições (total + por chave)."""
    try:
        with open(REQ_COUNTER_PATH, "r") as f:
            data = json.load(f)
            # Migrar formato antigo {total: N} para novo
            if "por_chave" not in data:
                data = {"total": data.get("total", 0), "por_chave": {}}
            return data
    except Exception:
        return {"total": 0, "por_chave": {}}

def _ler_req_count() -> int:
    """Lê contador persistente total de requisições."""
    return _ler_req_data().get("total", 0)

def _ler_req_count_chave(api_key: str) -> int:
    """Lê contador de requisições para uma chave de API específica."""
    data = _ler_req_data()
    # Usar hash parcial da chave como identificador (segurança)
    chave_id = api_key[-12:] if len(api_key) > 12 else api_key
    return data.get("por_chave", {}).get(chave_id, 0)

def _incrementar_req_count(api_key: str = "") -> int:
    """Incrementa e salva contador persistente (total + por chave). Retorna novo total."""
    data = _ler_req_data()
    data["total"] = data.get("total", 0) + 1
    if api_key:
        chave_id = api_key[-12:] if len(api_key) > 12 else api_key
        data.setdefault("por_chave", {})
        data["por_chave"][chave_id] = data["por_chave"].get(chave_id, 0) + 1
    try:
        with open(REQ_COUNTER_PATH, "w") as f:
            json.dump(data, f)
    except Exception:
        pass
    return data["total"]

MODELOS_DISPONIVEIS = {
    # ===== PAGOS (ordenados por popularidade) =====
    "💰 OpenAI GPT-4o Mini — $0.15/1M in | $0.60/1M out": "openai/gpt-4o-mini",
    "💰 DeepSeek V3.2 — $0.14/1M in | $0.28/1M out": "deepseek/deepseek-v3.2",
    "💰 Gemini 2.5 Flash Lite — $0.02/1M in | $0.10/1M out": "google/gemini-2.5-flash-lite",
    "💰 Llama 3.1 8B — $0.02/1M in | $0.05/1M out": "meta-llama/llama-3.1-8b-instruct",
    "💰 Mistral Nemo — $0.03/1M in | $0.03/1M out": "mistralai/mistral-nemo",
    "💰 OpenAI GPT-OSS 120B — $0.20/1M in | $0.60/1M out": "openai/gpt-oss-120b",
    "💰 OpenAI GPT-OSS 20B — $0.06/1M in | $0.18/1M out": "openai/gpt-oss-20b",
    "💰 Qwen 3.5 Flash — $0.05/1M in | $0.30/1M out": "qwen/qwen3.5-flash-02-23",
    "💰 Xiaomi MiMo V2 Flash — $0.04/1M in | $0.11/1M out": "xiaomi/mimo-v2-flash",
    # ===== GRÁTIS (ordenados por popularidade) =====
    "🆓 Google Gemma 4 26B A4B (grátis)": "google/gemma-4-26b-a4b-it:free",
    "🆓 Qwen 3.6 Plus (grátis)": "qwen/qwen3.6-plus:free",
    "🆓 NVIDIA Nemotron Super 120B (grátis)": "nvidia/nemotron-3-super-120b-a12b:free",
    "🆓 StepFun 3.5 Flash (grátis)": "stepfun/step-3.5-flash:free",
    "🆓 Arcee Trinity Large (grátis)": "arcee-ai/trinity-large-preview:free",
    "🆓 Z-AI GLM 4.5 Air (grátis)": "z-ai/glm-4.5-air:free",
}

# ============================================================
# INICIALIZAÇÃO DO SESSION STATE
# ============================================================
def _get_api_key():
    """Tenta obter a chave do OpenRouter de múltiplas fontes."""
    # 1. Variável de ambiente / .env
    key = os.environ.get("OPENROUTER_API_KEY", "")
    if key:
        return key
    # 2. st.secrets (Streamlit Cloud) — top level ou dentro de [openrouter]
    try:
        key = st.secrets.get("OPENROUTER_API_KEY", "")
        if key:
            return key
        key = st.secrets.get("openrouter", {}).get("OPENROUTER_API_KEY", "")
        if key:
            return key
    except Exception:
        pass
    return ""

_defaults = {
    "babilaca_messages": [],
    "babilaca_api_key": _get_api_key(),
    "babilaca_modelo": "openai/gpt-4o-mini",
    "babilaca_alertas": [],
    "babilaca_docs_gerados": [],
    "babilaca_preferencias": {},
    "babilaca_respostas_salvas": [],
    "babilaca_sessao_reqs": 0,
    "babilaca_sessao_custo": 0.0,
    "babilaca_sessao_tokens_in": 0,
    "babilaca_sessao_tokens_out": 0,
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ============================================================
# BASE JURÍDICA (RAG SIMPLIFICADO) — arquivo externo
# ============================================================
import importlib.util as _ilu
_spec = _ilu.spec_from_file_location(
    "base_juridica",
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "base_juridica.py"),
)
_mod = _ilu.module_from_spec(_spec)
_spec.loader.exec_module(_mod)
BASE_JURIDICA = _mod.BASE_JURIDICA

# --- Base RAG expandida (INs, Decreto 10.024, Acórdãos TCU) ---
_RAG_JSON_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "rag_knowledge_base.json")
RAG_KNOWLEDGE_BASE = []
if os.path.isfile(_RAG_JSON_PATH):
    with open(_RAG_JSON_PATH, "r", encoding="utf-8") as _f:
        RAG_KNOWLEDGE_BASE = json.load(_f)

# --- Texto integral da Lei 14.133/2021 (parser automático) ---
_spec_lei = _ilu.spec_from_file_location(
    "lei_parser",
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "lei_parser.py"),
)
_mod_lei = _ilu.module_from_spec(_spec_lei)
_spec_lei.loader.exec_module(_mod_lei)
ARTIGOS_LEI_INTEGRAL = _mod_lei.get_artigos_lei()


def _score_item(item: dict, tokens: set[str]) -> float:
    """Calcula score de relevância de um item da base jurídica."""
    kws = set(w.lower() for w in item["palavras_chave"])
    all_text = (item["conteudo"] + " " + item["titulo"]).lower()
    score = sum(1 for t in tokens if any(t in kw for kw in kws))
    score += sum(0.3 for t in tokens if t in all_text)
    return score


def buscar_base_juridica(pergunta: str, top_n: int = 15) -> list[dict]:
    """Busca os trechos mais relevantes da base jurídica para a pergunta.
    Retorna uma mistura equilibrada de Leis, INs, Decretos, Legislação Complementar e Acórdãos TCU.
    Fontes: base curada + texto integral Lei 14.133 + RAG (INs, Decreto 10.024, Acórdãos TCU)."""
    tokens = set(re.findall(r"\w+", pergunta.lower()))

    # 1) Pontuar itens da base curada (INs, acórdãos, resumos manuais)
    scored = []
    for item in BASE_JURIDICA:
        score = _score_item(item, tokens)
        if score > 0:
            scored.append((score, item))

    # 2) Pontuar artigos do texto integral da lei
    artigos_ids_curados = {
        item["artigo"] for item in BASE_JURIDICA if "Lei 14.133" in item.get("fonte", "")
    }
    for item in ARTIGOS_LEI_INTEGRAL:
        if item["artigo"] in artigos_ids_curados:
            continue
        score = _score_item(item, tokens)
        if score > 0:
            scored.append((score, item))

    # 3) Pontuar itens da base RAG expandida (INs, Decreto 10.024, Acórdãos TCU)
    for item in RAG_KNOWLEDGE_BASE:
        score = _score_item(item, tokens)
        if score > 0:
            scored.append((score, item))

    scored.sort(key=lambda x: x[0], reverse=True)

    # Garantir mix equilibrado: leis + decretos + INs + acórdãos
    leis = [s for s in scored if "Lei 14.133" in s[1]["fonte"]]
    decretos = [s for s in scored if "Decreto" in s[1]["fonte"]]
    ins = [s for s in scored if "IN " in s[1]["fonte"]]
    acordaos = [s for s in scored if "TCU" in s[1]["fonte"] or "Acord" in s[1]["fonte"]]
    outras = [
        s for s in scored
        if s not in leis and s not in decretos and s not in ins and s not in acordaos
    ]
    resultado = []
    max_lei = max(3, top_n // 4)
    max_decreto = max(2, top_n // 5)
    max_in = max(2, top_n // 5)
    max_tcu = max(3, top_n // 4)
    max_outras = top_n - max_lei - max_decreto - max_in - max_tcu
    resultado.extend(leis[:max_lei])
    resultado.extend(decretos[:max_decreto])
    resultado.extend(ins[:max_in])
    resultado.extend(acordaos[:max_tcu])
    resultado.extend(outras[:max(1, max_outras)])
    # Completar com restantes se faltou
    ids_usados = {id(s) for s in resultado}
    for s in scored:
        if len(resultado) >= top_n:
            break
        if id(s) not in ids_usados:
            resultado.append(s)
    resultado.sort(key=lambda x: x[0], reverse=True)
    return [s[1] for s in resultado[:top_n]]


# ============================================================
# FUNÇÕES DE API (TOOLS)
# ============================================================

def buscar_licitacoes(palavra_chave: str, pagina: int = 1) -> dict:
    """Busca licitações no PNCP."""
    try:
        hoje = datetime.now()
        inicio = (hoje - timedelta(days=90)).strftime("%Y%m%d")
        fim = hoje.strftime("%Y%m%d")
        resp = requests.get(
            "https://pncp.gov.br/api/consulta/v1/contratacoes/publicacao",
            params={
                "dataInicial": inicio,
                "dataFinal": fim,
                "pagina": pagina,
                "tamanhoPagina": 10,
            },
            timeout=15,
        )
        if resp.status_code == 200:
            data = resp.json()
            # Filtrar por palavra-chave localmente se a API não suportar
            if palavra_chave:
                kw = palavra_chave.lower()
                if isinstance(data, list):
                    data = [
                        i for i in data
                        if kw in json.dumps(i, ensure_ascii=False).lower()
                    ]
                elif isinstance(data, dict) and "data" in data:
                    data["data"] = [
                        i for i in data.get("data", [])
                        if kw in json.dumps(i, ensure_ascii=False).lower()
                    ]
            return {"sucesso": True, "dados": data, "fonte": "PNCP"}
        return {"sucesso": False, "erro": f"HTTP {resp.status_code}"}
    except Exception as e:
        return {"sucesso": False, "erro": str(e)}


def buscar_atas(item_codigo: str = "", descricao: str = "", pagina: int = 1) -> dict:
    """Busca atas de registro de preços no ComprasGov."""
    try:
        params = {"tamanhoPagina": 10, "pagina": pagina}
        if item_codigo:
            params["codigoItemCatalogoMaterial"] = item_codigo
        if descricao:
            params["descricaoItem"] = descricao
        resp = requests.get(
            "https://dadosabertos.compras.gov.br/modulo-arp/2_consultarARPItem",
            params=params,
            timeout=15,
        )
        if resp.status_code == 200:
            return {"sucesso": True, "dados": resp.json(), "fonte": "ComprasGov – Atas (ARP)"}
        return {"sucesso": False, "erro": f"HTTP {resp.status_code}"}
    except Exception as e:
        return {"sucesso": False, "erro": str(e)}


def consultar_fornecedor(cnpj: str) -> dict:
    """Consulta dados do fornecedor via OpenCNPJ + BrasilAPI."""
    cnpj_limpo = re.sub(r"[./-]", "", cnpj.strip())
    resultado: dict = {"cnpj": cnpj_limpo, "sucesso": False}
    # Tentativa 1 — OpenCNPJ
    try:
        resp = requests.get(f"https://api.opencnpj.org/{cnpj_limpo}", timeout=8)
        if resp.status_code == 200:
            resultado.update({"sucesso": True, "dados": resp.json(), "fonte": "OpenCNPJ"})
            return resultado
    except Exception:
        pass
    # Tentativa 2 — BrasilAPI
    try:
        resp = requests.get(
            f"https://brasilapi.com.br/api/cnpj/v1/{cnpj_limpo}", timeout=8
        )
        if resp.status_code == 200:
            resultado.update({"sucesso": True, "dados": resp.json(), "fonte": "BrasilAPI"})
            return resultado
    except Exception:
        pass
    resultado["erro"] = "Não foi possível consultar o CNPJ nas APIs disponíveis."
    return resultado


def buscar_fornecedores_cnae(cnae: str, pagina: int = 1) -> dict:
    """Busca fornecedores por CNAE no ComprasGov."""
    try:
        resp = requests.get(
            "https://dadosabertos.compras.gov.br/modulo-fornecedor/1_consultarFornecedor",
            params={"cnae": cnae, "tamanhoPagina": 10, "pagina": pagina},
            timeout=15,
        )
        if resp.status_code == 200:
            return {"sucesso": True, "dados": resp.json(), "fonte": "ComprasGov – Fornecedores"}
        return {"sucesso": False, "erro": f"HTTP {resp.status_code}"}
    except Exception as e:
        return {"sucesso": False, "erro": str(e)}


def buscar_precos_mercado(descricao: str) -> dict:
    """Busca preços de mercado via atas de registro de preços."""
    resultado = buscar_atas(descricao=descricao)
    if resultado.get("sucesso"):
        dados = resultado.get("dados", [])
        if isinstance(dados, dict):
            dados = dados.get("resultado", dados.get("data", []))
        precos = []
        for item in (dados if isinstance(dados, list) else []):
            preco = item.get("valorUnitario") or item.get("precoUnitario")
            if preco:
                precos.append(float(preco))
        if precos:
            return {
                "sucesso": True,
                "quantidade": len(precos),
                "menor": min(precos),
                "maior": max(precos),
                "media": sum(precos) / len(precos),
                "mediana": sorted(precos)[len(precos) // 2],
                "precos": precos[:20],
                "fonte": "ComprasGov – Atas (ARP)",
            }
    return {"sucesso": False, "erro": "Nenhum preço encontrado para o item."}


# ============================================================
# UTILIDADES — EXTRAÇÃO DE TEXTO
# ============================================================

def extrair_texto_pdf(arquivo) -> str:
    """Extrai texto de um arquivo PDF enviado pelo usuário."""
    if not PDF_SUPPORT:
        return "[PyPDF2 não instalado — execute: pip install PyPDF2]"
    try:
        reader = PyPDF2.PdfReader(arquivo)
        textos = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                textos.append(t)
        return "\n".join(textos) if textos else "[Não foi possível extrair texto do PDF]"
    except Exception as e:
        return f"[Erro ao ler PDF: {e}]"


def extrair_dados_excel(arquivo) -> tuple[str, pd.DataFrame | None]:
    """Extrai dados de um arquivo Excel."""
    try:
        df = pd.read_excel(arquivo, engine="openpyxl")
        resumo = f"Planilha com {len(df)} linhas e {len(df.columns)} colunas.\n"
        resumo += f"Colunas: {', '.join(df.columns.astype(str))}\n"
        resumo += f"Primeiras linhas:\n{df.head(10).to_string()}"
        return resumo, df
    except Exception as e:
        return f"[Erro ao ler Excel: {e}]", None


# ============================================================
# COMUNICAÇÃO COM IA (OPENROUTER)
# ============================================================

SYSTEM_PROMPT = """Você é **O Babilaca**, um assistente jurídico inteligente especializado em licitações e contratações públicas brasileiras, com foco na Lei 14.133/2021, Decreto 10.024/2019 (Pregão Eletrônico), Instruções Normativas (SEGES/ME nº 5/2017, 58/2022, 65/2021, 67/2021) e jurisprudência do TCU.

═══════════════════════════════════════════════════════════
REGRA FUNDAMENTAL — CITAÇÕES SOMENTE DO CONTEXTO FORNECIDO
═══════════════════════════════════════════════════════════
Você receberá abaixo um bloco chamado CONTEXTO JURÍDICO VERIFICADO contendo artigos de lei, instruções normativas e acórdãos do TCU. Estas são as ÚNICAS fontes que você pode citar com números específicos.

• CITE SOMENTE artigos, INs e acórdãos que estejam EXPLICITAMENTE listados no CONTEXTO JURÍDICO VERIFICADO abaixo.
• NUNCA invente, deduza ou complete números de artigos, INs ou acórdãos a partir de sua memória ou treinamento.
• Se o contexto fornecido contém "Art. 17" da Lei 14.133, cite "Art. 17". Se NÃO contém, NÃO cite.
• Se nenhuma fonte do contexto for relevante para a pergunta, diga: "Não disponho de referência verificada sobre esse ponto específico. Consulte o texto integral da Lei 14.133/2021 ou a jurisprudência do TCU."
• É PROIBIDO citar qualquer Acórdão TCU cujo número não apareça literalmente no contexto.
• É PROIBIDO citar qualquer IN cujo número não apareça literalmente no contexto.
• É PROIBIDO citar artigos da Lei 14.133 cujo número não apareça literalmente no contexto.

REGRAS DE QUALIDADE:
1. Estruture a resposta de forma clara, usando as fontes disponíveis no contexto.
2. Se o contexto contiver fontes de diferentes categorias (Lei, Decreto, IN, TCU), organize em camadas:
   - **Base legal**: artigo da Lei 14.133/2021 (se disponível no contexto)
   - **Regulamentação**: Decreto 10.024/2019 ou IN aplicável (se disponível no contexto)
   - **Jurisprudência TCU**: Acórdão relevante (se disponível no contexto)
3. Se uma camada não tiver fonte no contexto, omita-a — NÃO a preencha com informações inventadas.
4. Forneça os links que acompanham cada fonte no contexto (campo 🔗).
5. Responda sempre em português brasileiro.
6. Destaque trechos relevantes com **negrito** e use > citação para trechos literais.
7. Quando utilizar dados de APIs, indique a fonte.
8. No final da resposta, inclua uma seção "📚 **Fontes consultadas**" listando APENAS as fontes do contexto que você efetivamente utilizou, com seus links.
9. Para temas genéricos onde você pode responder com conhecimento geral (ex: conceitos, boas práticas), faça-o — mas NÃO atribua números específicos de artigos/acórdãos/INs que não estejam no contexto.

FORMATO DE CITAÇÃO (somente com fontes do contexto):
- Lei: "Conforme o Art. XX da Lei 14.133/2021..."
- IN: "De acordo com o Art. XX da IN SEGES/ME nº XX/XXXX..."
- TCU: "Nesse sentido, o TCU no Acórdão X.XXX/XXXX-Plenário firmou entendimento de que..."

AVISO: Você é uma ferramenta de apoio. Recomende sempre que o usuário confirme nas fontes oficiais.

{contexto_juridico}

{dados_api}"""


def _build_system_prompt(contexto_items: list[dict], dados_api: str = "") -> str:
    if contexto_items:
        partes = [
            "═══════════════════════════════════════════════",
            "CONTEXTO JURÍDICO VERIFICADO",
            "(Cite SOMENTE as fontes listadas abaixo)",
            "═══════════════════════════════════════════════",
        ]
        for i, item in enumerate(contexto_items, 1):
            fonte = item['fonte']
            artigo = item.get('artigo', '')
            titulo = item['titulo']
            ref = f"{fonte} — {artigo}" if artigo else fonte
            partes.append(
                f"[FONTE {i}] {ref}\n"
                f"Título: {titulo}\n"
                f"Conteúdo: {item['conteudo']}\n"
                f"🔗 Link: {item['link']}"
            )
        partes.append("═══════════════════════════════════════════════")
        partes.append(f"Total de fontes disponíveis: {len(contexto_items)}. NÃO cite nenhuma fonte fora desta lista.")
        ctx = "\n\n".join(partes)
    else:
        ctx = "(Nenhuma fonte jurídica disponível no contexto. Responda com conhecimento geral, SEM citar números específicos de artigos, INs ou acórdãos.)"
    api_section = f"DADOS DE APIs:\n{dados_api}" if dados_api else ""
    return SYSTEM_PROMPT.format(contexto_juridico=ctx, dados_api=api_section)


def chamar_ia(
    mensagens: list[dict],
    system_prompt: str,
    modelo: str | None = None,
    temperatura: float = 0.3,
) -> str:
    """Envia mensagens ao OpenRouter e retorna a resposta."""
    api_key = st.session_state.get("babilaca_api_key", "")
    if not api_key:
        return "⚠️ Chave de API não configurada. Acesse a aba **Configurações**."
    model = modelo or st.session_state.get("babilaca_modelo", "openai/gpt-4o-mini")
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://atacotada.streamlit.app",
        "X-Title": "AtaCotada - O Babilaca",
    }
    msgs = [{"role": "system", "content": system_prompt}] + mensagens

    # Tenta com 4096 tokens; se créditos insuficientes (402), reduz automaticamente
    for max_tk in (4096, 2048, 1024):
        payload = {
            "model": model,
            "messages": msgs,
            "temperature": temperatura,
            "max_tokens": max_tk,
        }
        try:
            resp = requests.post(
                OPENROUTER_URL, headers=headers, json=payload, timeout=60,
            )
            if resp.status_code == 200:
                data = resp.json()
                # Incrementar contador persistente de requisições
                _incrementar_req_count(api_key)
                # Rastrear uso da sessão (tokens e custo)
                usage = data.get("usage", {})
                st.session_state["babilaca_sessao_reqs"] = st.session_state.get("babilaca_sessao_reqs", 0) + 1
                st.session_state["babilaca_sessao_tokens_in"] = st.session_state.get("babilaca_sessao_tokens_in", 0) + (usage.get("prompt_tokens", 0) or 0)
                st.session_state["babilaca_sessao_tokens_out"] = st.session_state.get("babilaca_sessao_tokens_out", 0) + (usage.get("completion_tokens", 0) or 0)
                custo_req = float(data.get("usage", {}).get("total_cost", 0) or 0)
                if not custo_req:
                    # Estimar custo se a API não retornar
                    custo_req = ((usage.get("prompt_tokens", 0) or 0) * 0.15 + (usage.get("completion_tokens", 0) or 0) * 0.60) / 1_000_000
                st.session_state["babilaca_sessao_custo"] = st.session_state.get("babilaca_sessao_custo", 0.0) + custo_req
                return data["choices"][0]["message"]["content"]

            # --- Tratamento de erros conhecidos ---
            if resp.status_code == 402:
                # Créditos insuficientes — tentar com menos tokens
                try:
                    err = resp.json().get("error", {})
                    meta = err.get("metadata", {})
                except Exception:
                    err, meta = {}, {}
                if max_tk > 1024:
                    continue  # retry com menos tokens
                return (
                    "⚠️ **Créditos insuficientes** na conta OpenRouter.\n\n"
                    "O modelo selecionado é pago e sua conta não tem saldo suficiente. "
                    "Opções:\n"
                    "- Troque para um modelo **grátis** (marcados com ⚠️ no seletor)\n"
                    "- Adicione créditos em https://openrouter.ai/settings/credits"
                )

            if resp.status_code == 404:
                return (
                    f"⚠️ **Modelo indisponível**: `{model}`\n\n"
                    "Possíveis causas:\n"
                    "- Restrições de privacidade/política de dados na sua conta OpenRouter\n"
                    "- Modelo temporariamente fora do ar\n\n"
                    "Troque para outro modelo no seletor acima. "
                    "Se persistir, verifique suas configurações em https://openrouter.ai/settings/privacy"
                )

            if resp.status_code == 403:
                return (
                    "⚠️ **Limite diário atingido** para o modelo pago selecionado.\n\n"
                    "Sua chave de API alcançou o limite de uso do dia. "
                    "Por favor, selecione um **modelo gratuito** no seletor acima "
                    "(marcados com ⚠️) para continuar utilizando o Babilaca hoje.\n\n"
                    "Você também pode gerenciar seus limites em https://openrouter.ai/settings/keys"
                )

            if resp.status_code == 429:
                return (
                    "⚠️ **Limite de requisições atingido**. "
                    "Aguarde alguns segundos e tente novamente, ou troque para outro modelo."
                )

            return f"⚠️ Erro na API (HTTP {resp.status_code}): {resp.text[:300]}"

        except requests.Timeout:
            return "⚠️ Tempo limite da requisição excedido. Tente novamente."
        except Exception as e:
            return f"⚠️ Erro ao chamar a IA: {e}"

    return "⚠️ Não foi possível obter resposta da IA. Verifique sua conexão e créditos."


# ============================================================
# ROTEAMENTO DE INTENÇÃO
# ============================================================

# ---- Consultar saldo da conta OpenRouter ----
def consultar_saldo_openrouter() -> dict | None:
    """Consulta saldo e uso da conta via API /auth/key."""
    api_key = st.session_state.get("babilaca_api_key", "")
    if not api_key:
        return None
    try:
        resp = requests.get(
            "https://openrouter.ai/api/v1/auth/key",
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=8,
        )
        if resp.status_code == 200:
            return resp.json().get("data", {})
    except Exception:
        pass
    return None


def extrair_cnpj(texto: str) -> str | None:
    m = re.search(r"\d{2}\.?\d{3}\.?\d{3}/?\d{4}-?\d{2}", texto)
    return re.sub(r"[./-]", "", m.group()) if m else None


def analisar_intencao(msg: str) -> list[str]:
    """Determina quais APIs chamar com base na mensagem."""
    low = msg.lower()
    acoes = []
    if any(k in low for k in ["licitação", "licitacao", "pregão", "pregao", "edital", "contratação pública"]):
        acoes.append("buscar_licitacoes")
    if any(k in low for k in ["ata de registro", "ata ", "adesão", "adesao", "srp", "registro de preço"]):
        acoes.append("buscar_atas")
    if extrair_cnpj(msg):
        acoes.append("consultar_fornecedor")
    if any(k in low for k in ["fornecedor", "empresa", "cnae"]):
        if "consultar_fornecedor" not in acoes:
            acoes.append("buscar_fornecedores")
    if any(k in low for k in ["preço", "preco", "cotação", "cotacao", "valor de mercado", "quanto custa"]):
        acoes.append("buscar_precos")
    return acoes


def executar_apis(msg: str, acoes: list[str]) -> str:
    """Executa as APIs necessárias e retorna texto com resultados."""
    partes = []
    for acao in acoes:
        if acao == "buscar_licitacoes":
            palavras = re.sub(r"[^\w\s]", "", msg).strip()
            resultado = buscar_licitacoes(palavras)
            partes.append(f"[API — Licitações PNCP]\n{json.dumps(resultado, ensure_ascii=False, default=str)[:2000]}")
        elif acao == "buscar_atas":
            palavras = re.sub(r"[^\w\s]", "", msg).strip()
            resultado = buscar_atas(descricao=palavras)
            partes.append(f"[API — Atas ARP]\n{json.dumps(resultado, ensure_ascii=False, default=str)[:2000]}")
        elif acao == "consultar_fornecedor":
            cnpj = extrair_cnpj(msg)
            if cnpj:
                resultado = consultar_fornecedor(cnpj)
                partes.append(f"[API — Fornecedor {cnpj}]\n{json.dumps(resultado, ensure_ascii=False, default=str)[:2000]}")
        elif acao == "buscar_fornecedores":
            pass  # Requer CNAE específico; tratado via chat interativo
        elif acao == "buscar_precos":
            palavras = re.sub(r"[^\w\s]", "", msg).strip()
            resultado = buscar_precos_mercado(palavras)
            partes.append(f"[API — Preços de mercado]\n{json.dumps(resultado, ensure_ascii=False, default=str)[:2000]}")
    return "\n\n".join(partes)


# ============================================================
# GERAÇÃO DE DOCUMENTOS
# ============================================================

def _sanitize_for_pdf(text: str) -> str:
    """Remove/substitui caracteres que não são suportados pela fonte Helvetica (Latin-1)."""
    replacements = {
        "\u2013": "-", "\u2014": "-", "\u2015": "-",  # en-dash, em-dash
        "\u2018": "'", "\u2019": "'",  # aspas simples curvas
        "\u201c": '"', "\u201d": '"',  # aspas duplas curvas
        "\u2022": "-",  # bullet
        "\u2026": "...",  # ellipsis
        "\u00a0": " ",  # non-breaking space
        "\u200b": "",  # zero-width space
        "\u2010": "-", "\u2011": "-",  # hyphens
        "\u2212": "-",  # minus sign
        "\u2264": "<=", "\u2265": ">=",  # comparison
        "\u00b0": "o",  # degree -> o (nº)
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Remove emojis and other non-Latin-1 characters
    cleaned = []
    for ch in text:
        try:
            ch.encode("latin-1")
            cleaned.append(ch)
        except UnicodeEncodeError:
            cleaned.append(" ")
    return "".join(cleaned)


def gerar_pdf_documento(titulo: str, corpo: str) -> bytes:
    """Gera um PDF simples com título e corpo."""
    titulo = _sanitize_for_pdf(titulo)
    corpo = _sanitize_for_pdf(corpo)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)
    # Título
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(0, 26, 77)
    pdf.cell(0, 12, titulo, ln=True, align="C")
    pdf.ln(4)
    pdf.set_draw_color(212, 175, 55)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(8)
    # Data
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}", ln=True, align="R")
    pdf.ln(4)
    # Corpo
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(30, 30, 30)
    for linha in corpo.split("\n"):
        # Garante que X está na margem esquerda antes de cada renderização
        pdf.x = pdf.l_margin
        # Remove markdown bold/italic markers para o PDF
        linha_limpa = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", linha)
        linha_limpa = _sanitize_for_pdf(linha_limpa)
        if linha.startswith("## "):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 13)
            pdf.x = pdf.l_margin
            pdf.multi_cell(0, 7, _sanitize_for_pdf(linha.replace("## ", "")))
            pdf.set_font("Helvetica", "", 11)
        elif linha.startswith("### "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 11)
            pdf.x = pdf.l_margin
            pdf.multi_cell(0, 6, _sanitize_for_pdf(linha.replace("### ", "")))
            pdf.set_font("Helvetica", "", 11)
        elif linha.strip().startswith("- "):
            texto_bullet = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", linha.strip()[2:])
            pdf.x = pdf.l_margin
            pdf.multi_cell(0, 6, _sanitize_for_pdf(f"  - {texto_bullet}"))
        elif linha.strip() == "---":
            pdf.ln(3)
            pdf.set_draw_color(180, 180, 180)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(3)
        elif linha.strip() == "":
            pdf.ln(3)
        else:
            pdf.x = pdf.l_margin
            pdf.multi_cell(0, 6, linha_limpa)

    # Rodape
    pdf.ln(8)
    pdf.set_draw_color(212, 175, 55)
    pdf.set_line_width(0.5)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(3)
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 5, _sanitize_for_pdf("Gerado por AtaCotada - O Babilaca (IA) - Marinha do Brasil"), ln=True, align="C")

    return bytes(pdf.output())


TEMPLATE_DFD = """## DOCUMENTO DE FORMALIZAÇÃO DE DEMANDA (DFD)

### 1. Identificação do Requisitante
- Órgão/Entidade: {orgao}
- Setor Requisitante: {setor}
- Responsável: {responsavel}

### 2. Descrição da Necessidade
{necessidade}

### 3. Justificativa da Necessidade
{justificativa}

### 4. Quantidade Estimada
{quantidade}

### 5. Previsão de Data
- Data desejada para contratação: {data_prevista}

### 6. Grau de Prioridade
{prioridade}

### 7. Vinculação ao Plano de Contratações Anual (PCA)
{vinculacao_pca}

### 8. Base Legal
- Lei 14.133/2021, Art. 72, §1°
- IN SEGES/ME nº 58/2022

---
Documento gerado pelo sistema O Babilaca (IA) — Ferramenta de apoio.
Confirmar sempre nas fontes oficiais.
"""

TEMPLATE_TR = """## TERMO DE REFERÊNCIA

### 1. Objeto
{objeto}

### 2. Justificativa da Contratação
{justificativa}

### 3. Descrição Detalhada / Especificações
{especificacoes}

### 4. Quantidades e Unidades
{quantidades}

### 5. Valor Estimado
{valor_estimado}

### 6. Condições de Entrega / Execução
{condicoes_entrega}

### 7. Prazo de Execução / Vigência
{prazo}

### 8. Condições de Pagamento
{pagamento}

### 9. Critério de Julgamento
{criterio_julgamento}

### 10. Penalidades
Conforme previsto na Lei 14.133/2021, Arts. 155 a 163.

### 11. Base Legal
- Lei 14.133/2021
- IN SEGES/ME nº 65/2021

---
Documento gerado pelo sistema O Babilaca (IA) — Ferramenta de apoio.
Confirmar sempre nas fontes oficiais.
"""

TEMPLATE_JUSTIFICATIVA = """## JUSTIFICATIVA DE CONTRATAÇÃO

### 1. Contextualização
{contexto}

### 2. Necessidade da Contratação
{necessidade}

### 3. Análise de Mercado
{analise_mercado}

### 4. Justificativa do Preço
{justificativa_preco}

### 5. Modalidade Sugerida
{modalidade}

### 6. Fundamentação Legal
{fundamentacao}

### 7. Conclusão
{conclusao}

---
Documento gerado pelo sistema O Babilaca (IA) — Ferramenta de apoio.
Confirmar sempre nas fontes oficiais.
"""

TEMPLATE_ATA_RP = """## ATA DE REGISTRO DE PREÇOS

### PREÂMBULO
- **Processo Administrativo nº:** {processo}
- **Pregão Eletrônico nº:** {pregao}
- **Órgão Gerenciador:** {orgao}
- **CNPJ do Órgão:** {cnpj_orgao}
- **Vigência da Ata:** {vigencia}
- **Data de Assinatura:** {data_assinatura}

O(A) **{orgao}**, com sede em {endereco_orgao}, inscrito(a) no CNPJ sob nº {cnpj_orgao}, neste ato representado(a) pelo(a) {cargo_representante}, Sr(a). {representante}, nomeado(a) pela Portaria nº ____________, resolve REGISTRAR OS PREÇOS para eventual aquisição/contratação dos itens abaixo especificados, conforme condições estabelecidas no Edital do Pregão Eletrônico nº {pregao} e seus anexos.

---

### 1. DO OBJETO
Registro de preços para {objeto}, conforme especificações e condições estabelecidas no Termo de Referência (Anexo I do Edital).

### 2. DO FORNECEDOR REGISTRADO
- **Razão Social:** {fornecedor}
- **CNPJ:** {cnpj_fornecedor}
- **Endereço:** {endereco_fornecedor}
- **Representante Legal:** {representante_fornecedor}
- **Telefone/E-mail:** {contato_fornecedor}

### 3. DOS ITENS REGISTRADOS

| Item | Descrição | Unidade | Qtd Estimada | Valor Unitário (R$) | Valor Total (R$) |
|------|-----------|---------|-------------|--------------------|--------------------|
{tabela_itens}

**Valor Total Registrado:** {valor_total}

### 4. DA VALIDADE DA ATA
A presente Ata de Registro de Preços terá validade de **{vigencia}**, contados a partir da data de sua publicação no Diário Oficial, podendo ser prorrogada por igual período, nos termos do **Art. 84 da Lei 14.133/2021**.

### 5. DA REVISÃO E CANCELAMENTO DOS PREÇOS
5.1. Os preços registrados poderão ser revisados em decorrência de eventual redução dos preços praticados no mercado ou de fato que eleve o custo dos bens, obras ou serviços registrados, nos termos do **Art. 82, §5° da Lei 14.133/2021**.

5.2. O registro do preço do fornecedor será cancelado quando:
- Descumprir as condições da Ata de Registro de Preços;
- Não aceitar reduzir o preço registrado, na hipótese de este se tornar superior aos praticados no mercado;
- Sofrer sanção prevista nos incisos III ou IV do Art. 156 da Lei 14.133/2021.

### 6. DAS CONDIÇÕES DE ENTREGA / EXECUÇÃO
6.1. **Prazo de entrega:** {prazo_entrega}
6.2. **Local de entrega:** {local_entrega}
6.3. **Responsável pelo recebimento:** Comissão de Recebimento designada pelo Órgão Gerenciador.

### 7. DO PAGAMENTO
7.1. O pagamento será efetuado em até **{prazo_pagamento}** após o recebimento definitivo, mediante apresentação de nota fiscal/fatura.
7.2. A nota fiscal deverá ser acompanhada das certidões de regularidade fiscal e trabalhista vigentes.

### 8. DAS OBRIGAÇÕES DO FORNECEDOR
8.1. Entregar os materiais/prestar os serviços conforme especificações do Termo de Referência;
8.2. Manter, durante toda a vigência da Ata, as condições de habilitação exigidas na licitação;
8.3. Comunicar ao Órgão Gerenciador qualquer alteração de dados cadastrais;
8.4. Responsabilizar-se pelos encargos trabalhistas, previdenciários, fiscais e comerciais;
8.5. Substituir, às suas expensas, os materiais que apresentarem defeitos.

### 9. DAS OBRIGAÇÕES DO ÓRGÃO GERENCIADOR
9.1. Gerenciar a Ata de Registro de Preços;
9.2. Conduzir eventuais renegociações de preços;
9.3. Aplicar, garantida a ampla defesa e o contraditório, as penalidades cabíveis;
9.4. Providenciar a publicação da Ata e seus aditivos no Diário Oficial.

### 10. DA ADESÃO À ATA (CARONA)
10.1. A Ata de Registro de Preços poderá ser utilizada por órgãos e entidades não participantes, mediante prévia consulta ao Órgão Gerenciador e anuência do Fornecedor, nos termos do **Art. 86 da Lei 14.133/2021**.
10.2. O quantitativo para adesão não poderá exceder, por órgão/entidade, a **50% dos quantitativos** registrados na Ata.
10.3. O quantitativo total de adesões não poderá exceder, na totalidade, ao **dobro** do quantitativo registrado.

### 11. DAS PENALIDADES
Aplicam-se as sanções previstas nos **Arts. 155 a 163 da Lei 14.133/2021**, sem prejuízo das responsabilidades civil e criminal.

### 12. DA FUNDAMENTAÇÃO LEGAL
- Lei 14.133/2021, Arts. 82 a 86 (Registro de Preços)
- Decreto nº 11.462/2023 (regulamenta o SRP)
- Edital do Pregão Eletrônico nº {pregao} e seus anexos

### 13. DO FORO
Fica eleito o foro da Justiça Federal — Seção Judiciária de __________ para dirimir quaisquer questões oriundas desta Ata.

### 14. DAS DISPOSIÇÕES GERAIS
{observacoes}

---
Documento gerado pelo sistema O Babilaca (IA) — Ferramenta de apoio.
Confirmar sempre nas fontes oficiais.
"""

TEMPLATE_MINUTA_CONTRATO = """## MINUTA DE CONTRATO DE {tipo_contrato}

### PREÂMBULO
- **Processo Administrativo nº:** {processo}
- **Licitação nº:** {licitacao}
- **Modalidade:** {modalidade}
- **Contrato nº:** ____/{ano}
- **Vigência:** {vigencia}

O(A) **{orgao}** (CONTRATANTE), com sede em {endereco_orgao}, inscrito(a) no CNPJ sob nº {cnpj_orgao}, neste ato representado(a) pelo(a) {cargo_representante}, Sr(a). {representante}, e a empresa **{empresa}** (CONTRATADA), inscrita no CNPJ sob nº {cnpj_empresa}, com sede em {endereco_empresa}, representada pelo(a) Sr(a). {representante_empresa}, celebram o presente Contrato, nos termos da **Lei 14.133/2021** e demais legislações aplicáveis.

---

### CLÁUSULA PRIMEIRA — DO OBJETO
1.1. {objeto}
1.2. A CONTRATADA deverá fornecer/executar os itens/serviços conforme especificações do Termo de Referência (Anexo I).

### CLÁUSULA SEGUNDA — DOS ITENS / ESPECIFICAÇÕES

| Item | Descrição | Unidade | Quantidade | Valor Unitário (R$) | Valor Total (R$) |
|------|-----------|---------|-----------|--------------------|--------------------|
{tabela_itens}

**Valor Total do Contrato:** {valor_total}

### CLÁUSULA TERCEIRA — DO VALOR E DO PAGAMENTO
3.1. O valor total do presente contrato é de **{valor_total}** ({valor_extenso}).
3.2. O pagamento será efetuado em até **{prazo_pagamento}** após o recebimento definitivo dos materiais/serviços, mediante apresentação de nota fiscal/fatura devidamente atestada.
3.3. A nota fiscal deverá conter: descrição do item, quantidade, valor unitário e total, número do contrato e dados bancários da CONTRATADA.
3.4. Será retido na fonte os tributos previstos na legislação vigente (IR, CSLL, PIS, COFINS, ISS quando aplicável).

### CLÁUSULA QUARTA — DO PRAZO DE VIGÊNCIA E EXECUÇÃO
4.1. O presente contrato terá vigência de **{vigencia}**, contados a partir da data de sua assinatura.
4.2. O prazo de entrega/execução dos serviços será de **{prazo_execucao}**, contados a partir do recebimento da Ordem de Fornecimento/Serviço.
4.3. Os prazos poderão ser prorrogados nos termos do **Art. 107 da Lei 14.133/2021**.

### CLÁUSULA QUINTA — DA GARANTIA CONTRATUAL
5.1. A CONTRATADA deverá prestar garantia de execução do contrato no valor correspondente a **{percentual_garantia_contratual}** do valor total, em uma das seguintes modalidades (Art. 96, §1° da Lei 14.133/2021):
   - I — Caução em dinheiro ou em títulos da dívida pública;
   - II — Seguro-garantia;
   - III — Fiança bancária.
5.2. A garantia deverá ser apresentada no prazo de **10 (dez) dias úteis** após a assinatura do contrato.
5.3. A garantia somente será liberada após o integral cumprimento de todas as obrigações contratuais, inclusive quanto a multas aplicadas.

### CLÁUSULA SEXTA — DA GARANTIA DOS PRODUTOS / SERVIÇOS
6.1. Os materiais fornecidos deverão possuir garantia mínima de **{garantia_produto}**, contados a partir do recebimento definitivo.
6.2. Durante o período de garantia, a CONTRATADA deverá:
   - a) Substituir, sem ônus, os materiais que apresentarem defeitos de fabricação, vícios ou não conformidade com as especificações;
   - b) Realizar o reparo/substituição no prazo máximo de **{prazo_reparo}** após a notificação do defeito;
   - c) Arcar com todos os custos de transporte, frete, mão de obra e peças de reposição;
   - d) Fornecer assistência técnica autorizada pelo fabricante, quando aplicável.
6.3. A garantia não cobre defeitos decorrentes de mau uso, negligência, instalação inadequada ou caso fortuito e força maior devidamente comprovados.
6.4. Para serviços, a garantia cobre:
   - a) Vícios ou defeitos de execução que se manifestem dentro do prazo de garantia;
   - b) Conformidade com normas técnicas aplicáveis (ABNT, INMETRO e demais);
   - c) Reexecução integral do serviço defeituoso, às expensas da CONTRATADA.
6.5. O prazo de garantia será suspenso durante o período em que o material estiver em reparo ou substituição.
6.6. A recusa injustificada da CONTRATADA em atender chamados de garantia ensejará a aplicação das sanções previstas neste contrato.

### CLÁUSULA SÉTIMA — DAS OBRIGAÇÕES DA CONTRATADA
7.1. Executar o objeto conforme especificações do Termo de Referência e proposta apresentada;
7.2. Manter durante toda a execução do contrato as condições de habilitação exigidas na licitação;
7.3. Responsabilizar-se por todos os encargos trabalhistas, previdenciários, fiscais e comerciais;
7.4. Reparar, corrigir, remover ou substituir, às suas expensas, no total ou em parte, os materiais/serviços em que se verificarem vícios, defeitos ou incorreções;
7.5. Designar preposto para representá-la durante a execução contratual (Art. 118 da Lei 14.133/2021);
7.6. Comunicar ao CONTRATANTE qualquer anormalidade que interfira no cumprimento do contrato;
7.7. No caso de serviços continuados, apresentar mensalmente os comprovantes de regularidade fiscal, trabalhista e previdenciária;
7.8. Fornecer EPI e garantir condições de segurança do trabalho a seus empregados.

### CLÁUSULA OITAVA — DAS OBRIGAÇÕES DO CONTRATANTE
8.1. Proporcionar condições para a execução contratual;
8.2. Designar Fiscal e Gestor do Contrato (Art. 117 da Lei 14.133/2021);
8.3. Efetuar o pagamento conforme estipulado;
8.4. Comunicar à CONTRATADA as irregularidades observadas;
8.5. Aplicar as sanções cabíveis, garantidos o contraditório e a ampla defesa.

### CLÁUSULA NONA — DA FISCALIZAÇÃO
9.1. A fiscalização será exercida por servidor designado pelo CONTRATANTE, nos termos do **Art. 117 da Lei 14.133/2021**.
9.2. A fiscalização não exclui nem reduz a responsabilidade da CONTRATADA.

### CLÁUSULA DÉCIMA — DO RECEBIMENTO DO OBJETO
10.1. **Recebimento provisório:** em até **{prazo_recebimento_provisorio}** do término da entrega/execução, pelo fiscal do contrato.
10.2. **Recebimento definitivo:** em até **{prazo_recebimento_definitivo}** do recebimento provisório, pela comissão de recebimento ou gestor do contrato, após verificação da conformidade.
10.3. Nos casos de inconformidade, a CONTRATADA será notificada para correção em prazo razoável.

### CLÁUSULA DÉCIMA PRIMEIRA — DA ALTERAÇÃO CONTRATUAL
11.1. O contrato poderá ser alterado nos termos dos **Arts. 124 a 136 da Lei 14.133/2021**.
11.2. Acréscimos ou supressões de até **25%** do valor inicial atualizado (ou até 50% para reforma de edifícios/equipamentos).

### CLÁUSULA DÉCIMA SEGUNDA — DAS PENALIDADES
12.1. Pelo inadimplemento, aplicam-se as sanções dos **Arts. 155 a 163 da Lei 14.133/2021**:
   - I — Advertência;
   - II — Multa moratória de **{multa_moratoria}** por dia de atraso, limitada a **{limite_multa}** do valor do contrato;
   - III — Multa compensatória de até **{multa_compensatoria}** do valor total, por inexecução parcial ou total;
   - IV — Impedimento de licitar e contratar pelo prazo de até 3 anos;
   - V — Declaração de inidoneidade para licitar ou contratar.
12.2. As multas poderão ser descontadas da garantia, de pagamentos devidos ou cobradas judicialmente.

### CLÁUSULA DÉCIMA TERCEIRA — DA RESCISÃO
13.1. O contrato poderá ser rescindido nas hipóteses previstas nos **Arts. 137 a 139 da Lei 14.133/2021**.

### CLÁUSULA DÉCIMA QUARTA — DA FUNDAMENTAÇÃO LEGAL
- Lei 14.133/2021 (Lei de Licitações e Contratos)
- Decreto nº 11.462/2023 (quando aplicável — SRP)
- IN SEGES/ME nº 5/2017 ou IN SEGES/ME nº 58/2022 (conforme o caso)
- Edital da Licitação nº {licitacao} e seus anexos

### CLÁUSULA DÉCIMA QUINTA — DO FORO
Fica eleito o foro da Justiça Federal — Seção Judiciária de __________ para dirimir quaisquer questões oriundas deste contrato.

### CLÁUSULA DÉCIMA SEXTA — DAS DISPOSIÇÕES GERAIS
{observacoes}

---

**CONTRATANTE:**
{orgao}
Nome: {representante}
Cargo: {cargo_representante}

**CONTRATADA:**
{empresa}
Nome: {representante_empresa}

**TESTEMUNHAS:**
1. Nome: _________________ CPF: _________________
2. Nome: _________________ CPF: _________________

---
Documento gerado pelo sistema O Babilaca (IA) — Ferramenta de apoio.
Confirmar sempre nas fontes oficiais.
"""


def _fmt_brl(valor: float) -> str:
    """Formata um valor float em reais BR: R$ 1.234,56"""
    if not valor:
        return "-"
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def gerar_mapa_comparativo_pdf(itens: list[dict], observacoes: str = "") -> bytes:
    """Gera PDF de mapa comparativo de preços com quantidade, totais e observações."""
    pdf = FPDF("L")  # Paisagem
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(0, 26, 77)
    pdf.cell(0, 10, _sanitize_for_pdf("MAPA COMPARATIVO DE PRECOS"), ln=True, align="C")
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}", ln=True, align="R")
    pdf.ln(4)

    # Cabeçalho da tabela
    pdf.set_fill_color(0, 26, 77)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 8)
    colunas = ["Item", _sanitize_for_pdf("Descricao"), "Qtd", "Forn. 1", "Forn. 2", "Forn. 3", _sanitize_for_pdf("Media"), "Mediana", "Total Est."]
    larguras = [10, 55, 14, 32, 32, 32, 30, 30, 32]
    for col, larg in zip(colunas, larguras):
        pdf.cell(larg, 8, col, border=1, align="C", fill=True)
    pdf.ln()

    pdf.set_text_color(30, 30, 30)
    pdf.set_font("Helvetica", "", 8)
    total_geral = 0.0
    for i, item in enumerate(itens, 1):
        precos = [item.get("preco1", 0), item.get("preco2", 0), item.get("preco3", 0)]
        validos = [p for p in precos if p > 0]
        media = sum(validos) / len(validos) if validos else 0
        mediana = sorted(validos)[len(validos) // 2] if validos else 0
        qtd = item.get("quantidade", 1) or 1
        total_item = media * qtd
        total_geral += total_item
        pdf.cell(10, 7, str(i), border=1, align="C")
        pdf.cell(55, 7, _sanitize_for_pdf(str(item.get("descricao", ""))[:38]), border=1)
        pdf.cell(14, 7, str(int(qtd)), border=1, align="C")
        for p in precos:
            pdf.cell(32, 7, _fmt_brl(p), border=1, align="R")
        pdf.cell(30, 7, _fmt_brl(media), border=1, align="R")
        pdf.cell(30, 7, _fmt_brl(mediana), border=1, align="R")
        pdf.cell(32, 7, _fmt_brl(total_item), border=1, align="R")
        pdf.ln()

    # Linha de total geral
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_fill_color(240, 245, 255)
    largura_antes = sum(larguras[:-1])
    pdf.cell(largura_antes, 8, "TOTAL GERAL", border=1, align="R", fill=True)
    pdf.cell(larguras[-1], 8, _fmt_brl(total_geral), border=1, align="R", fill=True)
    pdf.ln()

    # Observações
    if observacoes and observacoes.strip():
        pdf.ln(6)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(0, 26, 77)
        pdf.cell(0, 7, _sanitize_for_pdf("OBSERVACOES"), ln=True)
        pdf.set_draw_color(212, 175, 55)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(3)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(30, 30, 30)
        pdf.multi_cell(0, 6, _sanitize_for_pdf(observacoes.strip()))

    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 5, _sanitize_for_pdf("Documento gerado pelo sistema O Babilaca (IA) - Ferramenta de apoio. Confirmar sempre nas fontes oficiais."), ln=True)
    return bytes(pdf.output())


def gerar_docx_documento(titulo: str, corpo: str) -> bytes:
    """Gera um documento Word (.docx) editável a partir de título e corpo markdown simples."""
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(30, 30, 30)

    # Título
    p_titulo = doc.add_heading(titulo, level=1)
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_titulo.runs:
        run.font.color.rgb = RGBColor(0, 26, 77)

    # Data
    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_data = p_data.add_run(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run_data.font.size = Pt(9)
    run_data.font.color.rgb = RGBColor(100, 100, 100)

    for linha in corpo.split("\n"):
        linha_limpa = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", linha)
        if linha.startswith("## "):
            doc.add_heading(linha.replace("## ", ""), level=2)
        elif linha.startswith("### "):
            doc.add_heading(linha.replace("### ", ""), level=3)
        elif linha.strip().startswith("- "):
            texto_bullet = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", linha.strip()[2:])
            doc.add_paragraph(texto_bullet, style="List Bullet")
        elif linha.strip() == "---":
            p = doc.add_paragraph()
            p.add_run("_" * 60).font.color.rgb = RGBColor(180, 180, 180)
        elif linha.strip() == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(linha_limpa)

    # Rodapé
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    run_f = p_footer.add_run("Gerado por AtaCotada - O Babilaca (IA) - Marinha do Brasil")
    run_f.font.size = Pt(7)
    run_f.font.italic = True
    run_f.font.color.rgb = RGBColor(150, 150, 150)
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def gerar_mapa_comparativo_xlsx(itens: list[dict], observacoes: str = "") -> bytes:
    """Gera planilha Excel editável do mapa comparativo de preços."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Mapa Comparativo"

    # Estilos
    header_font = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
    header_fill = PatternFill(start_color="001A4D", end_color="001A4D", fill_type="solid")
    total_fill = PatternFill(start_color="F0F5FF", end_color="F0F5FF", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    center = Alignment(horizontal="center", vertical="center")
    right_align = Alignment(horizontal="right", vertical="center")

    # Título
    ws.merge_cells("A1:I1")
    ws["A1"] = "MAPA COMPARATIVO DE PREÇOS"
    ws["A1"].font = Font(name="Calibri", bold=True, size=14, color="001A4D")
    ws["A1"].alignment = Alignment(horizontal="center")

    ws.merge_cells("A2:I2")
    ws["A2"] = f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws["A2"].font = Font(name="Calibri", size=9, color="666666")
    ws["A2"].alignment = Alignment(horizontal="right")

    # Cabeçalho
    headers = ["Item", "Descrição", "Qtd", "Fornecedor 1", "Fornecedor 2", "Fornecedor 3", "Média", "Mediana", "Total Estimado"]
    larguras = [6, 40, 8, 16, 16, 16, 16, 16, 18]
    for col_idx, (header, larg) in enumerate(zip(headers, larguras), 1):
        cell = ws.cell(row=4, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center
        ws.column_dimensions[get_column_letter(col_idx)].width = larg

    # Dados
    total_geral = 0.0
    for i, item in enumerate(itens, 1):
        row = i + 4
        precos = [item.get("preco1", 0), item.get("preco2", 0), item.get("preco3", 0)]
        validos = [p for p in precos if p > 0]
        media = sum(validos) / len(validos) if validos else 0
        mediana = sorted(validos)[len(validos) // 2] if validos else 0
        qtd = item.get("quantidade", 1) or 1
        total_item = media * qtd
        total_geral += total_item

        valores = [i, item.get("descricao", ""), int(qtd),
                    item.get("preco1", 0), item.get("preco2", 0), item.get("preco3", 0),
                    media, mediana, total_item]
        for col_idx, val in enumerate(valores, 1):
            cell = ws.cell(row=row, column=col_idx, value=val)
            cell.border = thin_border
            if col_idx >= 4:
                cell.number_format = '#,##0.00'
                cell.alignment = right_align
            elif col_idx in (1, 3):
                cell.alignment = center

    # Total geral
    total_row = len(itens) + 5
    ws.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=8)
    cell_label = ws.cell(row=total_row, column=1, value="TOTAL GERAL")
    cell_label.font = Font(name="Calibri", bold=True, size=10)
    cell_label.fill = total_fill
    cell_label.border = thin_border
    cell_label.alignment = Alignment(horizontal="right")
    cell_total = ws.cell(row=total_row, column=9, value=total_geral)
    cell_total.font = Font(name="Calibri", bold=True, size=10)
    cell_total.fill = total_fill
    cell_total.border = thin_border
    cell_total.number_format = '#,##0.00'
    cell_total.alignment = right_align

    # Observações
    if observacoes and observacoes.strip():
        obs_row = total_row + 2
        ws.cell(row=obs_row, column=1, value="OBSERVAÇÕES:").font = Font(name="Calibri", bold=True, size=10)
        ws.merge_cells(start_row=obs_row + 1, start_column=1, end_row=obs_row + 1, end_column=9)
        ws.cell(row=obs_row + 1, column=1, value=observacoes.strip())

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ============================================================
# ANÁLISE INTELIGENTE
# ============================================================

def analisar_sobrepreco(precos: list[float], valor_referencia: float) -> dict:
    """Detecta possível sobrepreço comparando com média/mediana."""
    if not precos:
        return {"analise": "Dados insuficientes para análise."}
    media = sum(precos) / len(precos)
    mediana = sorted(precos)[len(precos) // 2]
    desvio_media = ((valor_referencia - media) / media * 100) if media else 0
    desvio_mediana = ((valor_referencia - mediana) / mediana * 100) if mediana else 0
    if desvio_mediana > 30:
        risco = "🔴 ALTO"
        msg = "Valor significativamente acima da mediana de mercado."
    elif desvio_mediana > 15:
        risco = "🟡 MÉDIO"
        msg = "Valor acima da mediana. Recomenda-se justificativa detalhada."
    elif desvio_mediana > 0:
        risco = "🟢 BAIXO"
        msg = "Valor ligeiramente acima da mediana, dentro da faixa aceitável."
    else:
        risco = "🟢 OK"
        msg = "Valor igual ou abaixo da mediana de mercado."
    return {
        "risco": risco,
        "mensagem": msg,
        "media": media,
        "mediana": mediana,
        "desvio_media_pct": round(desvio_media, 2),
        "desvio_mediana_pct": round(desvio_mediana, 2),
        "valor_referencia": valor_referencia,
    }


def classificar_risco_juridico(descricao_cenario: str) -> str:
    """Usa IA para classificar o risco jurídico com mapa de riscos."""
    prompt_risco = (
        "Analise o cenário de contratação pública abaixo e produza:\n\n"
        "1. **CLASSIFICAÇÃO GERAL**: Risco BAIXO, MÉDIO ou ALTO com justificativa.\n\n"
        "2. **MAPA DE RISCOS** — Para CADA categoria abaixo, classifique Probabilidade "
        "(Baixa / Média / Alta) e Impacto (Baixo / Médio / Alto), e descreva brevemente o risco:\n"
        "   - Risco de Nulidade do Processo\n"
        "   - Risco de Responsabilização do Gestor\n"
        "   - Risco de Impugnação / Recurso\n"
        "   - Risco de Prejuízo ao Erário\n"
        "   - Risco Reputacional\n\n"
        "Use formato de tabela Markdown para o mapa de riscos com colunas: "
        "Categoria | Probabilidade | Impacto | Nível | Descrição\n\n"
        "3. **RECOMENDAÇÕES** de mitigação para os riscos identificados.\n\n"
        "Fundamente toda a análise na Lei 14.133/2021 e jurisprudência do TCU.\n\n"
        f"CENÁRIO: {descricao_cenario}"
    )
    ctx = buscar_base_juridica(descricao_cenario)
    system = _build_system_prompt(ctx)
    return chamar_ia([{"role": "user", "content": prompt_risco}], system)


# ============================================================
# CHECKLIST POR MODALIDADE
# ============================================================
CHECKLIST_MODALIDADES = {
    "Dispensa por Cotação Eletrônica": [
        "Documento de Formalização de Demanda (DFD)",
        "Estudo Técnico Preliminar (ETP) — quando aplicável",
        "Termo de Referência (TR)",
        "Enquadramento legal da dispensa (Art. 75, I ou II, Lei 14.133/2021)",
        "Pesquisa de Preços (mínimo 3 fontes — IN 65/2021)",
        "Mapa Comparativo de Preços",
        "Justificativa para contratação direta",
        "Aviso de Dispensa Eletrônica publicado no PNCP (Art. 75, §3°)",
        "Comprovante de abertura no sistema de cotação eletrônica (ComprasGov)",
        "Propostas recebidas no sistema eletrônico",
        "Ata / Relatório da sessão de dispensa eletrônica",
        "Razão da escolha do fornecedor vencedor",
        "Justificativa do preço",
        "Documentação de habilitação do fornecedor",
        "Comprovante de regularidade fiscal e trabalhista",
        "Parecer jurídico (Art. 72, Lei 14.133/2021)",
        "Autorização da autoridade competente",
        "Nota de Empenho",
        "Publicação do resultado no PNCP",
    ],
    "Pregão Eletrônico (PE)": [
        "Documento de Formalização de Demanda (DFD)",
        "Estudo Técnico Preliminar (ETP)",
        "Mapa de Riscos",
        "Termo de Referência (TR)",
        "Pesquisa de Preços (mínimo 3 fontes — IN 65/2021)",
        "Mapa Comparativo de Preços",
        "Justificativa da modalidade (pregão)",
        "Minuta do Edital",
        "Parecer jurídico sobre o Edital (Art. 53, §4°)",
        "Designação do pregoeiro e equipe de apoio",
        "Autorização da autoridade competente",
        "Publicação do Edital no PNCP e DOU",
        "Ata da sessão pública",
        "Relatório de adjudicação e homologação",
        "Comprovante de publicação do resultado no PNCP",
    ],
    "TJDL — Dispensa de Licitação": [
        "Documento de Formalização de Demanda (DFD)",
        "Termo de Justificativa de Dispensa de Licitação (TJDL)",
        "Enquadramento legal da dispensa (Art. 75, Lei 14.133/2021)",
        "Razão da escolha do fornecedor (Art. 72, III)",
        "Justificativa do preço (Art. 72, IV)",
        "Pesquisa de Preços (mínimo 3 fontes — IN 65/2021)",
        "Mapa Comparativo de Preços",
        "Proposta formal do fornecedor",
        "Documentação de habilitação do fornecedor",
        "Comprovante de regularidade fiscal e trabalhista",
        "Parecer jurídico (Art. 72, III, Lei 14.133/2021)",
        "Autorização da autoridade competente",
        "Nota de Empenho",
        "Publicação no PNCP (Art. 94, §1°)",
    ],
    "TJIL — Inexigibilidade de Licitação": [
        "Documento de Formalização de Demanda (DFD)",
        "Termo de Justificativa de Inexigibilidade de Licitação (TJIL)",
        "Enquadramento legal da inexigibilidade (Art. 74, Lei 14.133/2021)",
        "Comprovação de inviabilidade de competição",
        "Razão da escolha do fornecedor com justificativa de exclusividade ou notória especialização",
        "Justificativa do preço (compatibilidade com mercado)",
        "Pesquisa de Preços comprovando compatibilidade",
        "Proposta formal do fornecedor",
        "Atestado de exclusividade ou comprovação de notória especialização (quando aplicável)",
        "Documentação de habilitação do fornecedor",
        "Comprovante de regularidade fiscal e trabalhista",
        "Parecer jurídico (Art. 72, III, Lei 14.133/2021)",
        "Autorização da autoridade competente",
        "Nota de Empenho",
        "Publicação no PNCP (Art. 94, §1°)",
    ],
    "Adesão a Ata de Registro de Preços (Carona)": [
        "Documento de Formalização de Demanda (DFD)",
        "Estudo Técnico Preliminar (ETP)",
        "Justificativa da vantajosidade da adesão",
        "Pesquisa de Preços comprovando compatibilidade com o mercado",
        "Cópia da Ata de Registro de Preços vigente",
        "Cópia do Edital da licitação que gerou a Ata",
        "Ofício de consulta ao órgão gerenciador",
        "Autorização do órgão gerenciador",
        "Concordância do fornecedor",
        "Comprovante de regularidade fiscal e trabalhista do fornecedor",
        "Parecer jurídico",
        "Autorização da autoridade competente",
        "Publicação no PNCP",
    ],
    "Participação em IRP (Intenção de Registro de Preços)": [
        "Documento de Formalização de Demanda (DFD)",
        "Estudo Técnico Preliminar (ETP)",
        "Termo de Referência (TR) com especificações e quantitativos do órgão",
        "Pesquisa de Preços (mínimo 3 fontes — IN 65/2021)",
        "Justificativa da necessidade de participação na IRP",
        "Ofício de manifestação de interesse ao órgão gerenciador",
        "Comprovante de registro da participação no ComprasGov (dentro do prazo de 8 dias úteis)",
        "Planilha de quantitativos e especificações enviada ao gerenciador",
        "Confirmação de aceite pelo órgão gerenciador",
        "Autorização da autoridade competente para participar",
        "Parecer jurídico (quando exigido pelo órgão)",
        "Acompanhamento da sessão pública (registro de participante)",
        "Cópia da Ata de Registro de Preços assinada",
        "Nota de Empenho (quando da efetiva contratação)",
    ],
    "Renovação de Ata por Mais 1 Ano": [
        "Ata de Registro de Preços vigente (dentro do prazo de validade)",
        "Ofício de solicitação de prorrogação ao órgão gerenciador",
        "Manifestação de concordância do órgão gerenciador",
        "Concordância expressa do(s) fornecedor(es) beneficiário(s)",
        "Pesquisa de Preços atualizada comprovando vantajosidade",
        "Comprovação de que os preços registrados permanecem compatíveis com o mercado",
        "Justificativa técnica da necessidade de prorrogação",
        "Comprovante de regularidade fiscal e trabalhista atualizado do fornecedor",
        "Parecer jurídico sobre a prorrogação (Art. 84, §3°, Lei 14.133/2021)",
        "Autorização da autoridade competente",
        "Termo Aditivo ou Apostilamento de prorrogação da Ata",
        "Publicação da prorrogação no PNCP",
    ],
    "Termo Aditivo": [
        "Contrato administrativo vigente",
        "Justificativa técnica e/ou econômica para a alteração contratual",
        "Identificação da hipótese legal (Art. 124 ou Art. 125, Lei 14.133/2021)",
        "Documentação comprobatória da necessidade (planilhas, laudos, pareceres técnicos)",
        "Pesquisa de preços atualizada (quando houver acréscimo ou supressão de valores)",
        "Concordância formal do contratado (quando exigida)",
        "Comprovante de regularidade fiscal e trabalhista atualizado do contratado",
        "Demonstrativo do impacto financeiro (acréscimos/supressões dentro dos limites legais)",
        "Parecer jurídico sobre a legalidade do aditamento (Art. 53, Lei 14.133/2021)",
        "Autorização da autoridade competente",
        "Minuta do Termo Aditivo assinada pelas partes",
        "Publicação do extrato do Termo Aditivo no PNCP (Art. 94, Lei 14.133/2021)",
        "Atualização do registro no sistema de gestão de contratos",
    ],
    "Apostilamento": [
        "Contrato administrativo vigente",
        "Identificação do tipo de apostilamento (reajuste, repactuação ou alteração não substancial)",
        "Documento comprobatório do índice de reajuste (IPCA, INPC, IGP-M ou índice setorial previsto no contrato)",
        "Memória de cálculo do reajuste ou repactuação",
        "Planilha demonstrando os novos valores unitários e globais",
        "Comprovação de que o reajuste está previsto no contrato original e/ou edital",
        "Comprovante de que decorreu o período mínimo de 1 ano para reajuste (Art. 135, Lei 14.133/2021)",
        "Parecer jurídico (quando exigido pelo órgão)",
        "Autorização da autoridade competente",
        "Registro do apostilamento no verso/margem do contrato ou em documento apartado",
        "Publicação do apostilamento no PNCP",
        "Atualização dos valores no sistema de gestão de contratos",
    ],
}

# Informações explicativas de cada modalidade
INFO_MODALIDADES = {
    "Dispensa por Cotação Eletrônica": {
        "descricao": (
            "A **Dispensa por Cotação Eletrônica** é o procedimento de contratação direta "
            "realizado por meio de sistema eletrônico (ComprasGov), previsto no **Art. 75, §3°, "
            "da Lei 14.133/2021**. Substitui a antiga dispensa presencial, trazendo mais "
            "transparência e competitividade."
        ),
        "quando_aplicar": (
            "- **Bens e serviços comuns** com valor até **R$ 59.906,02** (Art. 75, II — valor atualizado pelo Decreto 12.343/2024);\n"
            "- **Serviços de engenharia** com valor até **R$ 119.812,07** (Art. 75, I);\n"
            "- Sempre que a dispensa se enquadrar nos incisos I ou II do Art. 75."
        ),
        "limites": (
            "| Tipo | Limite (atualizado) |\n"
            "|------|---------------------|\n"
            "| Bens e serviços comuns | R$ 59.906,02 |\n"
            "| Serviços de engenharia/obras | R$ 119.812,07 |\n\n"
            "*Valores atualizados pelo Decreto 12.343/2024. Atualização anual pelo IPCA.*"
        ),
        "como_montar": (
            "1. Elaborar DFD e TR com especificação clara do objeto;\n"
            "2. Realizar pesquisa de preços (IN 65/2021 — mín. 3 fontes);\n"
            "3. Publicar Aviso de Dispensa Eletrônica no PNCP (prazo mín. 3 dias úteis);\n"
            "4. Abrir sessão no ComprasGov para recebimento de propostas;\n"
            "5. Avaliar propostas e habilitação do vencedor;\n"
            "6. Adjudicar, obter parecer jurídico e autorizar;\n"
            "7. Emitir nota de empenho e publicar resultado no PNCP."
        ),
    },
    "Pregão Eletrônico (PE)": {
        "descricao": (
            "O **Pregão Eletrônico** é a modalidade obrigatória para aquisição de **bens e "
            "serviços comuns** (Art. 6°, XLI e Art. 28, I, Lei 14.133/2021). Caracteriza-se pela "
            "inversão de fases (proposta antes da habilitação) e lances em tempo real."
        ),
        "quando_aplicar": (
            "- Aquisição de **bens comuns** (qualquer valor acima do limite de dispensa);\n"
            "- Contratação de **serviços comuns** (inclusive de engenharia, quando aplicável);\n"
            "- Sempre que o objeto puder ser definido por especificações usuais de mercado."
        ),
        "limites": (
            "**Não possui limite máximo de valor.** O Pregão Eletrônico pode ser utilizado para "
            "qualquer valor, desde que o objeto seja bem ou serviço comum.\n\n"
            "*Limite mínimo implícito:* acima dos valores de dispensa (R$ 59.906,02 para bens/serviços)."
        ),
        "como_montar": (
            "1. Elaborar DFD, ETP e Mapa de Riscos;\n"
            "2. Elaborar Termo de Referência detalhado;\n"
            "3. Realizar pesquisa de preços (IN 65/2021);\n"
            "4. Elaborar minuta do Edital e obter parecer jurídico;\n"
            "5. Designar pregoeiro e equipe de apoio;\n"
            "6. Publicar Edital no PNCP e DOU (prazo mín. 8 dias úteis);\n"
            "7. Realizar sessão pública com lances no ComprasGov;\n"
            "8. Habilitar vencedor, adjudicar e homologar;\n"
            "9. Publicar resultado no PNCP."
        ),
    },
    "TJDL — Dispensa de Licitação": {
        "descricao": (
            "O **Termo de Justificativa de Dispensa de Licitação (TJDL)** é usado quando a "
            "administração precisa contratar **sem licitar**, em hipóteses taxativamente "
            "previstas no **Art. 75 da Lei 14.133/2021**, como valor baixo, emergência, "
            "ou situações específicas."
        ),
        "quando_aplicar": (
            "- Contratações de **baixo valor** (Art. 75, I e II) sem cotação eletrônica;\n"
            "- **Emergência ou calamidade** (Art. 75, VIII) — prazo máx. 1 ano;\n"
            "- **Guerra ou grave perturbação da ordem** (Art. 75, VII);\n"
            "- Aquisição de **materiais das Forças Armadas** (Art. 75, IX);\n"
            "- Demais hipóteses dos incisos III a XVI do Art. 75."
        ),
        "limites": (
            "| Hipótese | Limite |\n"
            "|----------|--------|\n"
            "| Obras/serviços de engenharia (Art. 75, I) | R$ 119.812,07 |\n"
            "| Bens/serviços comuns (Art. 75, II) | R$ 59.906,02 |\n"
            "| Emergência (Art. 75, VIII) | Sem limite de valor (prazo máx. 1 ano) |\n\n"
            "*Obs: No TJDL, a contratação é feita diretamente com fornecedor escolhido, sem cotação eletrônica.*"
        ),
        "como_montar": (
            "1. Elaborar DFD com justificativa da necessidade;\n"
            "2. Elaborar TJDL com enquadramento legal (inciso específico do Art. 75);\n"
            "3. Justificar razão da escolha do fornecedor e do preço;\n"
            "4. Realizar pesquisa de preços (IN 65/2021);\n"
            "5. Obter proposta formal e documentação do fornecedor;\n"
            "6. Obter parecer jurídico;\n"
            "7. Autorização da autoridade competente;\n"
            "8. Emitir nota de empenho e publicar no PNCP."
        ),
    },
    "TJIL — Inexigibilidade de Licitação": {
        "descricao": (
            "O **Termo de Justificativa de Inexigibilidade de Licitação (TJIL)** é utilizado "
            "quando há **inviabilidade de competição**, conforme **Art. 74 da Lei 14.133/2021**. "
            "Não se trata de escolher não licitar — a licitação é impossível."
        ),
        "quando_aplicar": (
            "- **Fornecedor exclusivo** comprovado (Art. 74, I);\n"
            "- Contratação de profissional de **notória especialização** (Art. 74, III);\n"
            "- Contratação de **artista consagrado** (Art. 74, II);\n"
            "- **Credenciamento** quando convém à administração contratar todos (Art. 74, IV);\n"
            "- Serviços de **treinamento** com profissional/empresa de notória especialização."
        ),
        "limites": (
            "**Não possui limite de valor.** A inexigibilidade pode ser utilizada para qualquer valor, "
            "desde que comprovada a inviabilidade de competição.\n\n"
            "*Atenção: a justificativa de preço é obrigatória mesmo sem limite, conforme Art. 72, IV.*"
        ),
        "como_montar": (
            "1. Elaborar DFD com descrição detalhada da necessidade;\n"
            "2. Elaborar TJIL demonstrando a inviabilidade de competição;\n"
            "3. Comprovar exclusividade (atestado) ou notória especialização;\n"
            "4. Justificar escolha do fornecedor e compatibilidade do preço;\n"
            "5. Obter proposta formal e habilitação do fornecedor;\n"
            "6. Obter parecer jurídico;\n"
            "7. Autorização da autoridade competente;\n"
            "8. Emitir nota de empenho e publicar no PNCP."
        ),
    },
    "Adesão a Ata de Registro de Preços (Carona)": {
        "descricao": (
            "A **Adesão a Ata de Registro de Preços (Carona)** permite que um órgão ou entidade "
            "que **não participou** da licitação original adira a uma Ata vigente de outro "
            "órgão, conforme **Art. 86 da Lei 14.133/2021**."
        ),
        "quando_aplicar": (
            "- Quando existe **Ata de Registro de Preços vigente** de outro órgão com o objeto desejado;\n"
            "- Quando a adesão for **mais vantajosa** que realizar novo processo;\n"
            "- Quando o órgão gerenciador e o fornecedor **concordarem** com a adesão;\n"
            "- Respeitado o limite de até **50% dos quantitativos** registrados na Ata (Art. 86, §4°)."
        ),
        "limites": (
            "| Regra | Limite |\n"
            "|-------|--------|\n"
            "| Quantitativo máximo por adesão | 50% do quantitativo registrado |\n"
            "| Quantitativo total de adesões | Até o dobro do quantitativo original |\n\n"
            "*O preço deve ser compatível com o praticado no mercado (pesquisa obrigatória).*"
        ),
        "como_montar": (
            "1. Elaborar DFD e ETP justificando a adesão;\n"
            "2. Demonstrar vantajosidade em relação a novo processo;\n"
            "3. Realizar pesquisa de preços para comprovar compatibilidade;\n"
            "4. Obter cópia da Ata e do Edital original;\n"
            "5. Solicitar autorização ao órgão gerenciador;\n"
            "6. Obter concordância do fornecedor;\n"
            "7. Verificar habilitação e regularidade fiscal do fornecedor;\n"
            "8. Obter parecer jurídico e autorização da autoridade;\n"
            "9. Publicar no PNCP."
        ),
    },
    "Participação em IRP (Intenção de Registro de Preços)": {
        "descricao": (
            "A **Participação em IRP** ocorre quando seu órgão identifica uma **Intenção de "
            "Registro de Preços** publicada por outro órgão (gerenciador) e decide aderir como "
            "**órgão participante**, informando seus quantitativos para serem incluídos na "
            "licitação, conforme **Art. 86, §1° da Lei 14.133/2021** e **Decreto 11.462/2023**. "
            "Diferente da Adesão (Carona), aqui o órgão participa **antes** da licitação."
        ),
        "quando_aplicar": (
            "- Quando o órgão identifica uma **IRP publicada no ComprasGov** com objeto compatível;\n"
            "- Quando há **demanda real** pelo bem ou serviço registrado;\n"
            "- Quando a participação é mais vantajosa que conduzir licitação própria;\n"
            "- Para **planejar aquisições futuras** com preços já registrados;\n"
            "- O prazo para manifestar interesse é de **mínimo 8 dias úteis** após publicação da IRP."
        ),
        "limites": (
            "**Não possui limite financeiro próprio.** O valor depende do quantitativo informado "
            "e dos preços que forem registrados na Ata.\n\n"
            "| Regra | Detalhe |\n"
            "|-------|--------|\n"
            "| Prazo para manifestar interesse | Mínimo 8 dias úteis (Art. 86, §1°) |\n"
            "| Quantitativo | Definido pelo órgão participante conforme sua demanda |\n"
            "| Vigência da Ata | Até 1 ano, prorrogável por mais 1 ano (Art. 84, III) |\n"
            "| Obrigação de contratar | Não obrigatória para participante (facultativa) |\n\n"
            "*Vantagem: o órgão participante tem prioridade sobre órgãos não participantes (carona).*"
        ),
        "como_montar": (
            "1. Monitorar as IRPs publicadas no **ComprasGov** que atendam à necessidade;\n"
            "2. Elaborar **DFD** justificando a demanda e a participação na IRP;\n"
            "3. Elaborar **ETP** e **TR** com especificações e quantitativos do órgão;\n"
            "4. Realizar pesquisa de preços para comprovar compatibilidade;\n"
            "5. Registrar manifestação de interesse no ComprasGov **dentro do prazo de 8 dias úteis**;\n"
            "6. Enviar planilha de quantitativos e especificações ao órgão gerenciador;\n"
            "7. Obter autorização da autoridade competente;\n"
            "8. Acompanhar a sessão pública da licitação;\n"
            "9. Após homologação, receber cópia da Ata de Registro de Preços;\n"
            "10. Quando necessitar, emitir nota de empenho para efetivar a contratação."
        ),
    },
    "Renovação de Ata por Mais 1 Ano": {
        "descricao": (
            "A **Renovação (prorrogação) de Ata de Registro de Preços** permite estender a "
            "vigência da Ata por mais **1 (um) ano**, conforme **Art. 84, §3° da Lei 14.133/2021**. "
            "A prorrogação depende de concordância do fornecedor e comprovação de que os preços "
            "continuam vantajosos em relação ao mercado."
        ),
        "quando_aplicar": (
            "- Quando a **Ata de Registro de Preços** está próxima do vencimento (vigência original de 1 ano);\n"
            "- Quando os **preços registrados continuam compatíveis** com o mercado;\n"
            "- Quando há **interesse da Administração** em manter o registro;\n"
            "- Quando o **fornecedor concorda** com a prorrogação nas mesmas condições;\n"
            "- Prorrogação permitida por **no máximo 1 ano** além da vigência original (Art. 84, §3°)."
        ),
        "limites": (
            "| Regra | Detalhe |\n"
            "|-------|--------|\n"
            "| Vigência original da Ata | Até 1 ano |\n"
            "| Prorrogação máxima | Mais 1 ano (total máximo: 2 anos) |\n"
            "| Condição obrigatória | Pesquisa demonstrando vantajosidade dos preços |\n"
            "| Concordância do fornecedor | Obrigatória |\n\n"
            "*A prorrogação não é automática — requer análise de vantajosidade e formalização.*"
        ),
        "como_montar": (
            "1. Verificar se a Ata ainda está dentro do prazo de vigência;\n"
            "2. Realizar **pesquisa de preços atualizada** para comprovar vantajosidade;\n"
            "3. Solicitar **concordância formal do fornecedor** para prorrogação;\n"
            "4. Verificar **regularidade fiscal e trabalhista** atualizada do fornecedor;\n"
            "5. Elaborar **justificativa técnica** da necessidade de prorrogação;\n"
            "6. Consultar o **órgão gerenciador** da Ata (se não for o próprio órgão);\n"
            "7. Obter **parecer jurídico** sobre a prorrogação;\n"
            "8. Obter **autorização da autoridade competente**;\n"
            "9. Formalizar via **Termo Aditivo** ou **Apostilamento**;\n"
            "10. **Publicar** a prorrogação no PNCP."
        ),
    },
    "Termo Aditivo": {
        "descricao": (
            "O **Termo Aditivo** é o instrumento formal utilizado para **modificar cláusulas contratuais** "
            "originalmente pactuadas, conforme **Arts. 124 a 136 da Lei 14.133/2021**. Diferente do "
            "apostilamento, o Termo Aditivo exige **acordo bilateral** (assinatura de ambas as partes) "
            "e é obrigatório quando a alteração afeta o equilíbrio econômico-financeiro ou modifica "
            "condições substanciais do contrato."
        ),
        "quando_aplicar": (
            "**Casos em que ocorre o Termo Aditivo (Art. 124, Lei 14.133/2021):**\n\n"
            "- **Acréscimo ou supressão** de quantitativos (até 25% do valor inicial, ou 50% para reforma — Art. 125);\n"
            "- **Prorrogação de prazo** de vigência ou execução do contrato;\n"
            "- **Alteração do objeto** por necessidade superveniente (modificação do projeto, especificações);\n"
            "- **Reequilíbrio econômico-financeiro** por fatos imprevisíveis (Art. 124, II, d);\n"
            "- **Substituição de garantia** contratual;\n"
            "- **Alteração da forma de pagamento** (nunca antecipando sem contraprestação);\n"
            "- **Inclusão ou exclusão de obrigações** das partes;\n"
            "- **Alteração do regime de execução** (quando justificado tecnicamente);\n"
            "- Qualquer modificação que altere **cláusulas essenciais** do contrato."
        ),
        "limites": (
            "| Regra | Limite |\n"
            "|-------|--------|\n"
            "| Acréscimo (bens/serviços/obras) | Até 25% do valor inicial atualizado (Art. 125) |\n"
            "| Acréscimo (reforma de edifício/equipamento) | Até 50% do valor inicial (Art. 125, §1°) |\n"
            "| Supressão | Até 25% do valor inicial atualizado |\n"
            "| Supressão acima de 25% | Apenas por acordo entre as partes (Art. 125, §2°) |\n"
            "| Prorrogação (serviços contínuos) | Até 5 anos, prorrogável por mais 5 (Art. 107) |\n"
            "| Prorrogação (outros contratos) | Conforme prazo máximo do Art. 105/106 |\n\n"
            "*O Termo Aditivo não pode descaracterizar o objeto original do contrato nem do edital que o originou.*"
        ),
        "como_montar": (
            "1. Identificar a **necessidade da alteração** e a hipótese legal (Art. 124 ou 125);\n"
            "2. Elaborar **justificativa técnica e/ou econômica** detalhada;\n"
            "3. Quando envolver valores, realizar **pesquisa de preços** atualizada;\n"
            "4. Verificar se a alteração está **dentro dos limites legais** (25% ou 50%);\n"
            "5. Obter **concordância formal do contratado** (quando necessário);\n"
            "6. Verificar **regularidade fiscal e trabalhista** do contratado;\n"
            "7. Elaborar **demonstrativo do impacto financeiro**;\n"
            "8. Obter **parecer jurídico** (Art. 53, Lei 14.133/2021);\n"
            "9. Obter **autorização da autoridade competente**;\n"
            "10. Elaborar e assinar a **minuta do Termo Aditivo** (ambas as partes);\n"
            "11. Publicar o **extrato do Termo Aditivo no PNCP** (Art. 94);\n"
            "12. Atualizar o registro no **sistema de gestão de contratos**."
        ),
    },
    "Apostilamento": {
        "descricao": (
            "O **Apostilamento** é o registro administrativo **unilateral** de alterações contratuais "
            "que **não modificam cláusulas essenciais** do contrato, conforme **Art. 136 da Lei 14.133/2021**. "
            "É mais simples que o Termo Aditivo — não exige assinatura do contratado nem parecer jurídico "
            "obrigatório. Utilizado principalmente para **reajustes de preços** previstos no contrato e "
            "**compensações/penalizações** financeiras previstas."
        ),
        "quando_aplicar": (
            "**Casos em que ocorre o Apostilamento (Art. 136, Lei 14.133/2021):**\n\n"
            "- **Reajuste de preços** previsto no contrato (por índice — IPCA, INPC, IGP-M ou setorial);\n"
            "- **Repactuação** de contratos de serviços contínuos com dedicação exclusiva de mão de obra;\n"
            "- **Compensações ou penalizações financeiras** por condições de pagamento;\n"
            "- **Empenho de dotações orçamentárias** complementares até o limite do contrato;\n"
            "- Alterações que **não modifiquem** o objeto, as condições de execução nem o equilíbrio contratual;\n\n"
            "**NÃO cabe apostilamento quando:**\n"
            "- Houver alteração de **cláusula essencial** (prazo, objeto, condições de execução);\n"
            "- For necessário **acréscimo ou supressão** de quantitativos;\n"
            "- A alteração exigir **concordância do contratado** → usar Termo Aditivo."
        ),
        "limites": (
            "| Regra | Detalhe |\n"
            "|-------|--------|\n"
            "| Natureza | Ato **unilateral** da Administração |\n"
            "| Assinatura do contratado | **Não necessária** |\n"
            "| Parecer jurídico | Não obrigatório (recomendável em repactuações complexas) |\n"
            "| Prazo mínimo para reajuste | 1 ano a partir da data-base (Art. 135) |\n"
            "| Índice de reajuste | O previsto no contrato/edital |\n"
            "| Forma de registro | Anotação no verso/margem do contrato ou documento apartado |\n\n"
            "*O apostilamento é mais ágil que o Termo Aditivo, mas limitado a alterações não substanciais.*"
        ),
        "como_montar": (
            "1. Identificar o **tipo de apostilamento** (reajuste por índice, repactuação, etc.);\n"
            "2. Verificar se o **reajuste está previsto** no contrato/edital e qual o **índice aplicável**;\n"
            "3. Confirmar que transcorreu o **período mínimo de 1 ano** da data-base (Art. 135);\n"
            "4. Elaborar **memória de cálculo** detalhada com o índice acumulado;\n"
            "5. Elaborar **planilha** com os novos valores unitários e globais;\n"
            "6. Obter **parecer jurídico** quando exigido pelo órgão ou em casos complexos;\n"
            "7. Obter **autorização da autoridade competente**;\n"
            "8. Registrar o apostilamento no **verso/margem do contrato** ou em documento apartado;\n"
            "9. **Publicar** o apostilamento no PNCP;\n"
            "10. **Atualizar** os valores no sistema de gestão de contratos e empenho."
        ),
    },
}


def gerar_checklist_pdf(modalidade: str, itens: list[str], status: dict[str, bool]) -> bytes:
    """Gera um PDF bonito com o checklist do processo."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=25)
    pw = pdf.w - 40  # largura útil

    # === CABEÇALHO ===
    pdf.set_fill_color(0, 26, 77)
    pdf.rect(0, 0, 210, 50, "F")
    pdf.set_draw_color(212, 175, 55)
    pdf.set_line_width(1.2)
    pdf.line(0, 50, 210, 50)

    pdf.set_y(8)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(212, 175, 55)
    pdf.cell(0, 5, "MARINHA DO BRASIL", ln=True, align="C")
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(200, 200, 220)
    pdf.cell(0, 4, "Centro de Operações do Abastecimento", ln=True, align="C")
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 9, _sanitize_for_pdf("CHECKLIST DE PROCESSO"), ln=True, align="C")
    pdf.ln(1)

    # === FAIXA DOURADA COM MODALIDADE ===
    pdf.set_fill_color(212, 175, 55)
    pdf.rect(0, 52, 210, 14, "F")
    pdf.set_y(54)
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(0, 26, 77)
    pdf.cell(0, 9, _sanitize_for_pdf(modalidade), ln=True, align="C")
    pdf.ln(6)

    # === DATA ===
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}", ln=True, align="R")
    pdf.ln(2)

    # === LINHA SEPARADORA ===
    pdf.set_draw_color(212, 175, 55)
    pdf.set_line_width(0.5)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(6)

    # === LEGENDA ===
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 5, "[X] = Documento presente    [  ] = Documento pendente", ln=True, align="C")
    pdf.ln(6)

    # === ITENS DO CHECKLIST ===
    pdf.set_font("Helvetica", "", 11)
    for i, item in enumerate(itens, 1):
        checked = status.get(item, False)
        marca = "X" if checked else "  "

        y_antes = pdf.get_y()

        # Fundo alternado
        if i % 2 == 0:
            pdf.set_fill_color(240, 245, 255)
            pdf.rect(18, y_antes - 1, pw + 4, 9, "F")

        # Checkbox
        pdf.set_draw_color(0, 26, 77)
        pdf.set_line_width(0.4)
        pdf.rect(20, y_antes, 6, 6)
        if checked:
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(0, 120, 0)
            pdf.set_xy(20.5, y_antes - 0.5)
            pdf.cell(6, 7, "X", align="C")

        # Texto do item
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(30, 30, 30)
        pdf.set_xy(29, y_antes)
        pdf.multi_cell(pw - 12, 6, _sanitize_for_pdf(f"{i}. {item}"))
        pdf.ln(2)

    # === RODAPÉ - LINHA DOURADA ===
    pdf.ln(8)
    pdf.set_draw_color(212, 175, 55)
    pdf.set_line_width(0.8)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(4)

    # Resumo
    total = len(itens)
    marcados = sum(1 for it in itens if status.get(it, False))
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(0, 26, 77)
    pdf.cell(0, 7, f"Total: {marcados}/{total} documentos presentes", ln=True, align="C")
    pdf.ln(2)

    if marcados == total:
        pdf.set_text_color(0, 120, 0)
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 7, _sanitize_for_pdf("PROCESSO COMPLETO"), ln=True, align="C")
    else:
        pdf.set_text_color(200, 50, 50)
        pdf.set_font("Helvetica", "B", 12)
        faltam = total - marcados
        pdf.cell(0, 7, _sanitize_for_pdf(f"ATENÇÃO: {faltam} documento(s) pendente(s)"), ln=True, align="C")

    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 5, _sanitize_for_pdf("Gerado por AtaCotada — O Babilaca (IA) — Marinha do Brasil"), ln=True, align="C")

    return bytes(pdf.output())

def verificar_alertas(palavras_chave: list[str] | None = None) -> list[dict]:
    """Verifica alertas proativos consultando APIs."""
    alertas = []
    # 1. Novas licitações
    if palavras_chave:
        for kw in palavras_chave[:3]:
            res = buscar_licitacoes(kw)
            if res.get("sucesso"):
                dados = res["dados"]
                qtd = len(dados) if isinstance(dados, list) else len(dados.get("data", [])) if isinstance(dados, dict) else 0
                if qtd > 0:
                    alertas.append({
                        "tipo": "info",
                        "titulo": f"Novas licitações: '{kw}'",
                        "descricao": f"Encontradas {qtd} licitação(ões) nos últimos 90 dias.",
                        "data": datetime.now().isoformat(),
                    })
    # 2. Alerta padrão de legislação
    alertas.append({
        "tipo": "info",
        "titulo": "Atualização legislativa",
        "descricao": (
            "Verifique regularmente o Portal Nacional de Contratações Públicas (PNCP) "
            "e o DOU para atualizações na regulamentação da Lei 14.133/2021."
        ),
        "data": datetime.now().isoformat(),
    })
    return alertas


# ============================================================
# SISTEMA DE MEMÓRIA
# ============================================================

def carregar_memoria() -> dict:
    if os.path.exists(MEMORIA_PATH):
        try:
            with open(MEMORIA_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"preferencias": {}, "respostas_salvas": [], "alertas_palavras": []}


def salvar_memoria(dados: dict):
    try:
        with open(MEMORIA_PATH, "w", encoding="utf-8") as f:
            json.dump(dados, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def salvar_resposta(pergunta: str, resposta: str):
    mem = carregar_memoria()
    mem["respostas_salvas"].append({
        "pergunta": pergunta,
        "resposta": resposta,
        "data": datetime.now().isoformat(),
    })
    salvar_memoria(mem)


# ============================================================
# SIDEBAR — Navegação customizada (padrão do projeto)
# ============================================================
_acanto_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "Projeto Adesões", "acanto.png")
if os.path.exists(_acanto_path):
    with open(_acanto_path, "rb") as _f:
        _acanto_b64 = base64.b64encode(_f.read()).decode()
else:
    _acanto_b64 = None

with st.sidebar:
    if _acanto_b64:
        st.markdown(f'<div style="text-align:center;padding:1rem 0 0.5rem 0;"><img src="data:image/png;base64,{_acanto_b64}" style="max-width:70%;height:auto;"></div>', unsafe_allow_html=True)
    st.markdown("## MENU")
    st.markdown("---")
    st.page_link("streamlit_app.py", label="Cotação", icon="⚓")
    st.page_link("pages/Detalhes_Compra.py", label="Detalhes Compra", icon="🔍")
    st.page_link("pages/Adesões.py", label="Adesões", icon="🤝")
    st.page_link("pages/Notas_Fiscais.py", label="Notas Fiscais", icon="📄")
    st.page_link("pages/Banco_de_Fornecedores.py", label="Fornecedores", icon="🏢")
    st.page_link("pages/Consulta.py", label="Consulta CNPJ", icon="💻")
    st.page_link("pages/Web_Scraping.py", label="Web Scraping", icon="🕷️")
    st.page_link("pages/O_Babilaca_(IA).py", label="O Babilaca (IA)", icon="🧠")
    st.markdown("---")
    st.markdown("## LINKS ÚTEIS")
    st.markdown("""
    <div style="margin-bottom: 1rem;">
        <a href="https://freitasnascimentomarinha-debug.github.io/ShootMail/" target="_blank" style="color: #cbd5e1; text-decoration: none; font-size: 0.9rem; display: flex; align-items: center; gap: 0.5rem;">
            📧 Disparador de Emails
        </a>
    </div>
    <div style="margin-bottom: 1rem;">
        <a href="https://detetive-obtencao.vercel.app/" target="_blank" style="color: #cbd5e1; text-decoration: none; font-size: 0.9rem; display: flex; align-items: center; gap: 0.5rem;">
            🚨 Detetive Obtenção
        </a>
    </div>
    """, unsafe_allow_html=True)
    st.markdown('<div style="text-align:center;color:#d4af37;font-size:10px;font-weight:600;padding:0.3rem 0;white-space:nowrap;">Centro de Operações do Abastecimento</div>', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-footer">Marinha do Brasil<br>AtaCotada v1.0</div>', unsafe_allow_html=True)

# ============================================================
# HEADER
# ============================================================
st.markdown("""
<div class="babilaca-header">
    <h1>🧠 O Babilaca (IA)</h1>
    <p>Assistente Inteligente para Licitações e Contratações Públicas — Lei 14.133/2021</p>
</div>
<div class="disclaimer-box">
    ⚠️ Ferramenta de apoio. Confirmar sempre nas fontes oficiais.
</div>
""", unsafe_allow_html=True)

# ============================================================
# ABAS PRINCIPAIS
# ============================================================
tab_chat, tab_docs, tab_checklist, tab_ipca = st.tabs([
    "💬 Chat IA",
    "📄 Gerar Documentos com IA",
    "📋 Checklist de Processo",
    "📊 Cálculo IPCA",
])

# ============================================================
# TAB 1 — CHAT IA
# ============================================================
with tab_chat:
    st.markdown("### 💬 Converse com o Babilaca")

    # Contador de sessão
    _s_reqs = st.session_state.get("babilaca_sessao_reqs", 0)
    _s_custo = st.session_state.get("babilaca_sessao_custo", 0.0)
    _s_tk_in = st.session_state.get("babilaca_sessao_tokens_in", 0)
    _s_tk_out = st.session_state.get("babilaca_sessao_tokens_out", 0)
    _cor_custo = "#22c55e" if _s_custo < 0.01 else "#f59e0b" if _s_custo < 0.10 else "#ef4444"
    st.markdown(f"""
    <div style="background:rgba(10,22,40,0.7);border:1px solid #1e3a5f;border-radius:8px;padding:0.4rem 1rem;margin-bottom:0.6rem;display:flex;align-items:center;gap:1rem;flex-wrap:wrap;">
        <span style="color:#d4af37;font-weight:bold;font-size:0.82rem;">Selecione seu Modelo Treinado para Licitações</span>
        <span style="color:#60a5fa;font-size:0.75rem;font-weight:bold;">Sessão: {_s_reqs} req</span>
        <span style="color:{_cor_custo};font-size:0.75rem;font-weight:bold;">Gasto: ${_s_custo:.6f}</span>
        <span style="color:#94a3b8;font-size:0.72rem;">Tokens: {_s_tk_in:,} in | {_s_tk_out:,} out</span>
    </div>
    """, unsafe_allow_html=True)

    # Barra de ferramentas: Modelo | Limpar conversa | Anexar arquivo
    tb_col1, tb_col2, tb_col3 = st.columns([3, 1.3, 1.3])

    with tb_col1:
        modelo_nome = st.selectbox(
            "Modelo de IA",
            list(MODELOS_DISPONIVEIS.keys()),
            index=list(MODELOS_DISPONIVEIS.values()).index(st.session_state["babilaca_modelo"])
            if st.session_state["babilaca_modelo"] in MODELOS_DISPONIVEIS.values()
            else 0,
            key="toolbar_modelo",
            label_visibility="collapsed",
        )
        st.session_state["babilaca_modelo"] = MODELOS_DISPONIVEIS[modelo_nome]
    with tb_col2:
        if st.button("🗑️ Limpar conversa", use_container_width=True):
            st.session_state["babilaca_messages"] = []
            st.session_state.pop("_arquivo_texto", None)
            st.session_state.pop("_arquivo_nome", None)
            st.rerun()
    with tb_col3:
        arquivo_chat = st.file_uploader(
            "📎 Anexar",
            type=["pdf", "xlsx", "xls", "csv"],
            key="chat_upload",
            label_visibility="collapsed",
        )

    texto_arquivo = ""
    if arquivo_chat:
        if arquivo_chat.name.endswith(".pdf"):
            texto_arquivo = extrair_texto_pdf(arquivo_chat)
            st.success(f"PDF processado: {len(texto_arquivo)} caracteres extraídos.")
        elif arquivo_chat.name.endswith((".xlsx", ".xls")):
            texto_arquivo, _ = extrair_dados_excel(arquivo_chat)
            st.success("Planilha processada.")
        elif arquivo_chat.name.endswith(".csv"):
            try:
                df_csv = pd.read_csv(arquivo_chat)
                texto_arquivo = f"CSV com {len(df_csv)} linhas.\nColunas: {', '.join(df_csv.columns)}\n{df_csv.head(10).to_string()}"
                st.success("CSV processado.")
            except Exception as e:
                texto_arquivo = f"[Erro ao ler CSV: {e}]"
        # Persistir no session_state para não perder no rerun
        if texto_arquivo:
            st.session_state["_arquivo_texto"] = texto_arquivo
            st.session_state["_arquivo_nome"] = arquivo_chat.name
    else:
        # Recuperar texto previamente extraído (se houver)
        texto_arquivo = st.session_state.get("_arquivo_texto", "")

    # Histórico de mensagens
    for idx, msg in enumerate(st.session_state["babilaca_messages"]):
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg["role"] == "user":
                if st.button("🔄 Reenviar", key=f"resend_{idx}"):
                    st.session_state["_reenviar_prompt"] = msg["content"]
                    st.rerun()

    # Auto-scroll para manter sempre no final da conversa
    if st.session_state["babilaca_messages"]:
        st.markdown(
            """<div id="chat-anchor"></div>
            <script>
            (function() {
                const anchor = window.parent.document.getElementById('chat-anchor');
                if (anchor) anchor.scrollIntoView({behavior: 'smooth', block: 'end'});
                const chatContainer = window.parent.document.querySelector('section[data-testid="stChatMessageContainer"]');
                if (chatContainer) chatContainer.scrollTop = chatContainer.scrollHeight;
            })();
            </script>""",
            unsafe_allow_html=True,
        )

    # Processar reenvio de mensagem (botão 🔄)
    _reenviar = st.session_state.pop("_reenviar_prompt", None)

    # Input do usuário
    if prompt := (_reenviar or st.chat_input("Digite sua pergunta sobre licitações, legislação ou contratações...")):
        # Adicionar mensagem do usuário
        st.session_state["babilaca_messages"].append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("🧠 Pensando..."):
                # 1. Buscar contexto jurídico
                contexto = buscar_base_juridica(prompt)

                # 2. Analisar intenção e chamar APIs
                acoes = analisar_intencao(prompt)
                dados_api = ""
                if acoes:
                    dados_api = executar_apis(prompt, acoes)

                # 3. Contexto de arquivo (se houver)
                if texto_arquivo:
                    dados_api += f"\n\n[ARQUIVO ENVIADO PELO USUÁRIO]\n{texto_arquivo[:3000]}"

                # 4. Montar system prompt
                system = _build_system_prompt(contexto, dados_api)

                # 5. Preparar histórico (últimas 20 mensagens)
                historico = [
                    {"role": m["role"], "content": m["content"]}
                    for m in st.session_state["babilaca_messages"][-20:]
                ]

                # 6. Chamar IA
                resposta = chamar_ia(historico, system)

            st.markdown(resposta)

            # Botão salvar resposta (joinha)
            if st.button("👍", key=f"save_{len(st.session_state['babilaca_messages'])}"):
                salvar_resposta(prompt, resposta)
                st.toast("Resposta salva!")

        st.session_state["babilaca_messages"].append({"role": "assistant", "content": resposta})
        # Auto-scroll para o final da conversa
        st.markdown(
            """<script>
            const chatContainer = window.parent.document.querySelector('section[data-testid="stChatMessageContainer"]');
            if (chatContainer) { chatContainer.scrollTop = chatContainer.scrollHeight; }
            // Fallback: scroll the main container
            const main = window.parent.document.querySelector('.main');
            if (main) { main.scrollTop = main.scrollHeight; }
            </script>""",
            unsafe_allow_html=True,
        )

# ============================================================
# TAB 2 — DOCUMENTOS
# ============================================================
with tab_docs:
    st.markdown("### 📄 Gerar Documentos com IA")
    st.markdown("Gere documentos administrativos para facilitar a montagem do processo de licitação.")
    st.markdown("---")
    tipo_doc = st.selectbox(
        "Selecione o tipo de documento",
        ["DFD — Documento de Formalização de Demanda",
         "Termo de Referência",
         "Justificativa de Contratação",
         "Mapa Comparativo de Preços",
         "Cronograma Físico-Financeiro",
         "Memória de Cálculo",
         "Ata de Registro de Preços",
         "Minuta de Contrato de Serviço / Material",
         "Classificação de Risco Jurídico"],
    )

    # ---- DFD ----
    if "DFD" in tipo_doc:
        st.markdown("#### 📋 Gerar DFD")
        modo_dfd = st.radio("Modo de preenchimento", ["Manual", "Assistido por IA"], horizontal=True, key="modo_dfd", index=1)

        if modo_dfd == "Manual":
            with st.form("form_dfd"):
                c1, c2 = st.columns(2)
                with c1:
                    orgao = st.text_input("Órgão/Entidade")
                    setor = st.text_input("Setor Requisitante")
                    responsavel = st.text_input("Responsável")
                with c2:
                    data_prevista = st.date_input("Data prevista")
                    prioridade = st.selectbox("Prioridade", ["Alta", "Média", "Baixa"])
                necessidade = st.text_area("Descrição da necessidade", height=100)
                justificativa = st.text_area("Justificativa", height=100)
                quantidade = st.text_input("Quantidade estimada")
                vinculacao = st.text_input("Vinculação ao PCA (se houver)")
                submit_dfd = st.form_submit_button("📄 Gerar DFD", use_container_width=True)

            if submit_dfd:
                corpo = TEMPLATE_DFD.format(
                    orgao=orgao, setor=setor, responsavel=responsavel,
                    necessidade=necessidade, justificativa=justificativa,
                    quantidade=quantidade, data_prevista=data_prevista.strftime("%d/%m/%Y"),
                    prioridade=prioridade, vinculacao_pca=vinculacao or "Não informado",
                )
                st.markdown(corpo)
                pdf_bytes = gerar_pdf_documento("DOCUMENTO DE FORMALIZAÇÃO DE DEMANDA (DFD)", corpo)
                docx_bytes = gerar_docx_documento("DOCUMENTO DE FORMALIZAÇÃO DE DEMANDA (DFD)", corpo)
                _c_pdf, _c_docx = st.columns(2)
                with _c_pdf:
                    st.download_button("⬇️ Baixar DFD em PDF", pdf_bytes, "dfd.pdf", "application/pdf")
                with _c_docx:
                    st.download_button("📝 Baixar DFD em Word", docx_bytes, "dfd.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            desc_ia = st.text_area(
                "Descreva brevemente a necessidade de contratação e o Babilaca preencherá o DFD:",
                height=120,
                key="dfd_ia_input",
            )
            if st.button("🤖 Gerar DFD com IA", key="btn_dfd_ia"):
                if desc_ia:
                    with st.spinner("Gerando DFD com IA..."):
                        prompt_dfd = (
                            f"Gere um Documento de Formalização de Demanda (DFD) completo, "
                            f"seguindo a IN SEGES/ME nº 58/2022, para a seguinte necessidade:\n\n"
                            f"{desc_ia}\n\n"
                            f"Use o formato com seções: Identificação, Necessidade, Justificativa, "
                            f"Quantidade, Previsão de Data, Prioridade, Vinculação ao PCA."
                        )
                        ctx = buscar_base_juridica("DFD formalização demanda planejamento")
                        system = _build_system_prompt(ctx)
                        resultado = chamar_ia([{"role": "user", "content": prompt_dfd}], system)
                    st.markdown(resultado)
                    pdf_bytes = gerar_pdf_documento("DOCUMENTO DE FORMALIZAÇÃO DE DEMANDA (DFD)", resultado)
                    docx_bytes = gerar_docx_documento("DOCUMENTO DE FORMALIZAÇÃO DE DEMANDA (DFD)", resultado)
                    _c_pdf, _c_docx = st.columns(2)
                    with _c_pdf:
                        st.download_button("⬇️ Baixar DFD em PDF", pdf_bytes, "dfd_ia.pdf", "application/pdf")
                    with _c_docx:
                        st.download_button("📝 Baixar DFD em Word", docx_bytes, "dfd_ia.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.warning("Descreva a necessidade antes de gerar.")

    # ---- TERMO DE REFERÊNCIA ----
    elif "Termo de Referência" in tipo_doc:
        st.markdown("#### 📑 Gerar Termo de Referência")
        modo_tr = st.radio("Modo de preenchimento", ["Manual", "Assistido por IA"], horizontal=True, key="modo_tr", index=1)

        if modo_tr == "Manual":
            with st.form("form_tr"):
                objeto = st.text_area("Objeto da contratação", height=80)
                justificativa = st.text_area("Justificativa", height=80)
                especificacoes = st.text_area("Especificações técnicas", height=100)
                quantidades = st.text_input("Quantidades e unidades")
                valor_estimado = st.text_input("Valor estimado (R$)")
                condicoes = st.text_area("Condições de entrega/execução", height=60)
                prazo = st.text_input("Prazo de execução/vigência")
                pagamento = st.text_input("Condições de pagamento")
                criterio = st.selectbox("Critério de julgamento",
                    ["Menor preço", "Melhor técnica", "Técnica e preço", "Maior desconto", "Maior retorno econômico"])
                submit_tr = st.form_submit_button("📄 Gerar Termo de Referência", use_container_width=True)

            if submit_tr:
                corpo = TEMPLATE_TR.format(
                    objeto=objeto, justificativa=justificativa, especificacoes=especificacoes,
                    quantidades=quantidades, valor_estimado=valor_estimado,
                    condicoes_entrega=condicoes, prazo=prazo, pagamento=pagamento,
                    criterio_julgamento=criterio,
                )
                st.markdown(corpo)
                pdf_bytes = gerar_pdf_documento("TERMO DE REFERÊNCIA", corpo)
                docx_bytes = gerar_docx_documento("TERMO DE REFERÊNCIA", corpo)
                _c_pdf, _c_docx = st.columns(2)
                with _c_pdf:
                    st.download_button("⬇️ Baixar TR em PDF", pdf_bytes, "termo_referencia.pdf", "application/pdf")
                with _c_docx:
                    st.download_button("📝 Baixar TR em Word", docx_bytes, "termo_referencia.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            desc_tr = st.text_area(
                "Descreva o que precisa contratar e o Babilaca elaborará o Termo de Referência:",
                height=120,
                key="tr_ia_input",
            )
            arquivo_tr = st.file_uploader("Ou envie um arquivo com especificações (PDF/Excel)", type=["pdf", "xlsx"], key="tr_upload")
            if st.button("🤖 Gerar TR com IA", key="btn_tr_ia"):
                ctx_extra = ""
                if arquivo_tr:
                    if arquivo_tr.name.endswith(".pdf"):
                        ctx_extra = extrair_texto_pdf(arquivo_tr)
                    else:
                        ctx_extra, _ = extrair_dados_excel(arquivo_tr)
                if desc_tr or ctx_extra:
                    with st.spinner("Gerando Termo de Referência..."):
                        prompt_tr = (
                            f"Gere um Termo de Referência completo conforme a Lei 14.133/2021 "
                            f"para a seguinte contratação:\n\n{desc_tr}\n\n"
                        )
                        if ctx_extra:
                            prompt_tr += f"Dados adicionais do arquivo:\n{ctx_extra[:3000]}\n\n"
                        prompt_tr += (
                            "Inclua: Objeto, Justificativa, Especificações, Quantidades, "
                            "Valor Estimado, Condições de Entrega, Prazo, Pagamento, "
                            "Critério de Julgamento, Penalidades, Base Legal."
                        )
                        ctx = buscar_base_juridica("termo de referência contratação especificação")
                        system = _build_system_prompt(ctx)
                        resultado = chamar_ia([{"role": "user", "content": prompt_tr}], system)
                    st.markdown(resultado)
                    pdf_bytes = gerar_pdf_documento("TERMO DE REFERÊNCIA", resultado)
                    docx_bytes = gerar_docx_documento("TERMO DE REFERÊNCIA", resultado)
                    _c_pdf, _c_docx = st.columns(2)
                    with _c_pdf:
                        st.download_button("⬇️ Baixar TR em PDF", pdf_bytes, "tr_ia.pdf", "application/pdf")
                    with _c_docx:
                        st.download_button("📝 Baixar TR em Word", docx_bytes, "tr_ia.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.warning("Forneça uma descrição ou envie um arquivo.")

    # ---- JUSTIFICATIVA ----
    elif "Justificativa" in tipo_doc:
        st.markdown("#### 📝 Gerar Justificativa de Contratação")
        modo_just = st.radio("Modo", ["Manual", "Assistido por IA"], horizontal=True, key="modo_just", index=1)

        if modo_just == "Manual":
            with st.form("form_just"):
                contexto = st.text_area("Contextualização", height=80)
                necessidade = st.text_area("Necessidade da contratação", height=80)
                analise = st.text_area("Análise de mercado", height=80)
                just_preco = st.text_area("Justificativa do preço", height=60)
                modalidade = st.selectbox("Modalidade sugerida",
                    ["Pregão Eletrônico", "Concorrência", "Dispensa de Licitação",
                     "Inexigibilidade", "Adesão a Ata de Registro de Preços"])
                fundamentacao = st.text_area("Fundamentação legal", height=60)
                conclusao = st.text_area("Conclusão", height=60)
                submit_j = st.form_submit_button("📄 Gerar Justificativa", use_container_width=True)

            if submit_j:
                corpo = TEMPLATE_JUSTIFICATIVA.format(
                    contexto=contexto, necessidade=necessidade, analise_mercado=analise,
                    justificativa_preco=just_preco, modalidade=modalidade,
                    fundamentacao=fundamentacao, conclusao=conclusao,
                )
                st.markdown(corpo)
                pdf_bytes = gerar_pdf_documento("JUSTIFICATIVA DE CONTRATAÇÃO", corpo)
                docx_bytes = gerar_docx_documento("JUSTIFICATIVA DE CONTRATAÇÃO", corpo)
                _c_pdf, _c_docx = st.columns(2)
                with _c_pdf:
                    st.download_button("⬇️ Baixar Justificativa em PDF", pdf_bytes, "justificativa.pdf", "application/pdf")
                with _c_docx:
                    st.download_button("📝 Baixar Justificativa em Word", docx_bytes, "justificativa.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            desc_just = st.text_area(
                "Descreva a contratação que precisa justificar:",
                height=120,
                key="just_ia_input",
            )
            if st.button("🤖 Gerar Justificativa com IA", key="btn_just_ia"):
                if desc_just:
                    with st.spinner("Gerando justificativa..."):
                        prompt_just = (
                            f"Gere uma Justificativa de Contratação completa baseada na Lei 14.133/2021 "
                            f"para o seguinte cenário:\n\n{desc_just}\n\n"
                            f"Inclua: Contextualização, Necessidade, Análise de Mercado, "
                            f"Justificativa de Preço, Modalidade Sugerida, Fundamentação Legal, Conclusão."
                        )
                        ctx = buscar_base_juridica("justificativa contratação dispensa modalidade")
                        system = _build_system_prompt(ctx)
                        resultado = chamar_ia([{"role": "user", "content": prompt_just}], system)
                    st.markdown(resultado)
                    pdf_bytes = gerar_pdf_documento("JUSTIFICATIVA DE CONTRATAÇÃO", resultado)
                    docx_bytes = gerar_docx_documento("JUSTIFICATIVA DE CONTRATAÇÃO", resultado)
                    _c_pdf, _c_docx = st.columns(2)
                    with _c_pdf:
                        st.download_button("⬇️ Baixar em PDF", pdf_bytes, "justificativa_ia.pdf", "application/pdf")
                    with _c_docx:
                        st.download_button("📝 Baixar em Word", docx_bytes, "justificativa_ia.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.warning("Descreva o cenário.")

    # ---- MAPA COMPARATIVO ----
    elif "Mapa Comparativo" in tipo_doc:
        st.markdown("#### 📊 Mapa Comparativo de Preços")
        modo_mapa = st.radio("Modo", ["Manual", "Via Arquivo"], horizontal=True, key="modo_mapa")

        if modo_mapa == "Manual":
            n_itens = st.number_input("Número de itens", min_value=1, max_value=50, value=3, key="n_itens_mapa")
            itens_mapa = []
            for i in range(int(n_itens)):
                st.markdown(f"**Item {i + 1}**")
                c1, c2, c3, c4, c5 = st.columns([3, 1, 1.5, 1.5, 1.5])
                with c1:
                    desc = st.text_input("Descrição", key=f"mapa_desc_{i}")
                with c2:
                    qtd = st.number_input("Qtd", min_value=1, value=1, key=f"mapa_qtd_{i}")
                with c3:
                    p1 = st.number_input("Fornecedor 1 (R$)", min_value=0.0, format="%.2f", key=f"mapa_p1_{i}")
                with c4:
                    p2 = st.number_input("Fornecedor 2 (R$)", min_value=0.0, format="%.2f", key=f"mapa_p2_{i}")
                with c5:
                    p3 = st.number_input("Fornecedor 3 (R$)", min_value=0.0, format="%.2f", key=f"mapa_p3_{i}")
                itens_mapa.append({"descricao": desc, "quantidade": qtd, "preco1": p1, "preco2": p2, "preco3": p3})

            obs_mapa = st.text_area(
                "Observações (aparecerá abaixo da tabela no PDF)",
                height=80,
                key="mapa_obs",
                placeholder="Ex: Preços coletados em março/2026. Frete incluso...",
            )

            if st.button("📊 Gerar Mapa Comparativo", key="btn_mapa"):
                pdf_bytes = gerar_mapa_comparativo_pdf(itens_mapa, observacoes=obs_mapa)
                xlsx_bytes = gerar_mapa_comparativo_xlsx(itens_mapa, observacoes=obs_mapa)
                _c_pdf, _c_xlsx = st.columns(2)
                with _c_pdf:
                    st.download_button("⬇️ Baixar Mapa em PDF", pdf_bytes, "mapa_comparativo.pdf", "application/pdf")
                with _c_xlsx:
                    st.download_button("📊 Baixar Mapa em Excel", xlsx_bytes, "mapa_comparativo.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

                # Mostrar tabela na tela
                rows = []
                total_geral = 0.0
                for i, item in enumerate(itens_mapa, 1):
                    precos = [item["preco1"], item["preco2"], item["preco3"]]
                    validos = [p for p in precos if p > 0]
                    media = sum(validos) / len(validos) if validos else 0
                    q = item.get("quantidade", 1) or 1
                    total_item = media * q
                    total_geral += total_item
                    rows.append({
                        "Item": i,
                        "Descrição": item["descricao"],
                        "Qtd": int(q),
                        "Forn. 1": f"R$ {item['preco1']:,.2f}",
                        "Forn. 2": f"R$ {item['preco2']:,.2f}",
                        "Forn. 3": f"R$ {item['preco3']:,.2f}",
                        "Média": f"R$ {media:,.2f}" if validos else "-",
                        "Total Est.": f"R$ {total_item:,.2f}" if total_item else "-",
                    })
                df_mapa = pd.DataFrame(rows)
                st.dataframe(df_mapa, use_container_width=True, hide_index=True)
                st.markdown(f"**Total Geral Estimado: R$ {total_geral:,.2f}**")
                if obs_mapa and obs_mapa.strip():
                    st.info(f"**Observações:** {obs_mapa}")
        else:
            arq_mapa = st.file_uploader("Envie planilha ou PDF com itens e preços", type=["xlsx", "xls", "csv", "pdf"], key="mapa_upload")
            obs_mapa_arq = st.text_area(
                "Observações (aparecerá abaixo da tabela no PDF)",
                height=80,
                key="mapa_obs_arq",
                placeholder="Ex: Preços coletados em março/2026. Frete incluso...",
            )
            if arq_mapa:
                try:
                    if arq_mapa.name.lower().endswith(".pdf"):
                        texto_pdf = extrair_texto_pdf(arq_mapa)
                        st.text_area("Conteúdo extraído do PDF", texto_pdf[:3000], height=200, disabled=True)
                        st.info("PDF carregado. Use o modo Manual para preencher os dados com base no conteúdo extraído.")
                    elif arq_mapa.name.endswith(".csv"):
                        df_m = pd.read_csv(arq_mapa)
                    else:
                        df_m = pd.read_excel(arq_mapa)

                    if not arq_mapa.name.lower().endswith(".pdf"):
                        st.dataframe(df_m, use_container_width=True)
                        st.info("Revise os dados acima. O sistema tentará identificar colunas de descrição e preços.")
                        if st.button("📊 Gerar Mapa a partir da planilha", key="btn_mapa_arq"):
                            itens_arq = []
                            desc_col = None
                            for c in df_m.columns:
                                if any(k in c.lower() for k in ["desc", "item", "produto", "material", "objeto"]):
                                    desc_col = c
                                    break
                            if not desc_col:
                                desc_col = df_m.columns[0]
                            preco_cols = [c for c in df_m.columns if any(k in c.lower() for k in ["preco", "preço", "valor", "forn", "cota"])]
                            if not preco_cols:
                                preco_cols = [c for c in df_m.columns if df_m[c].dtype in ["float64", "int64"]]
                            # Tentar encontrar coluna de quantidade
                            qtd_col = None
                            for c in df_m.columns:
                                if any(k in c.lower() for k in ["qtd", "quant", "quantidade", "qty"]):
                                    qtd_col = c
                                    break
                            for _, row in df_m.iterrows():
                                item = {"descricao": str(row[desc_col])}
                                if qtd_col:
                                    try:
                                        item["quantidade"] = int(float(row[qtd_col]))
                                    except (ValueError, TypeError):
                                        item["quantidade"] = 1
                                else:
                                    item["quantidade"] = 1
                                for j, pc in enumerate(preco_cols[:3], 1):
                                    try:
                                        item[f"preco{j}"] = float(row[pc])
                                    except (ValueError, TypeError):
                                        item[f"preco{j}"] = 0
                                for j in range(len(preco_cols) + 1, 4):
                                    item[f"preco{j}"] = 0
                                itens_arq.append(item)
                            if itens_arq:
                                pdf_bytes = gerar_mapa_comparativo_pdf(itens_arq, observacoes=obs_mapa_arq)
                                xlsx_bytes = gerar_mapa_comparativo_xlsx(itens_arq, observacoes=obs_mapa_arq)
                                _c_pdf, _c_xlsx = st.columns(2)
                                with _c_pdf:
                                    st.download_button("⬇️ Baixar Mapa em PDF", pdf_bytes, "mapa_comparativo.pdf", "application/pdf")
                                with _c_xlsx:
                                    st.download_button("📊 Baixar Mapa em Excel", xlsx_bytes, "mapa_comparativo.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                except Exception as e:
                    st.error(f"Erro ao processar arquivo: {e}")

    # ---- CRONOGRAMA FÍSICO-FINANCEIRO ----
    elif "Cronograma" in tipo_doc:
        st.markdown("#### 📅 Cronograma Físico-Financeiro")
        st.markdown(
            "Gere um cronograma detalhado de execução e desembolso financeiro para obras, "
            "serviços de engenharia ou contratos de grande porte, conforme **Art. 25, §1° e "
            "Art. 46, Lei 14.133/2021**."
        )
        modo_crono = st.radio("Modo de preenchimento", ["Manual", "Assistido por IA"], horizontal=True, key="modo_crono", index=1)

        if modo_crono == "Manual":
            with st.form("form_cronograma"):
                st.markdown("**Dados Gerais**")
                c1, c2 = st.columns(2)
                with c1:
                    crono_objeto = st.text_area("Objeto da contratação", height=60, key="crono_obj")
                    crono_contrato = st.text_input("Nº do Contrato / Processo", key="crono_contrato")
                    crono_contratada = st.text_input("Empresa contratada", key="crono_contratada")
                with c2:
                    crono_valor_global = st.number_input("Valor Global do Contrato (R$)", min_value=0.0, format="%.2f", key="crono_valor")
                    crono_inicio = st.date_input("Data de início prevista", key="crono_inicio")
                    crono_prazo_meses = st.number_input("Prazo total (meses)", min_value=1, max_value=120, value=12, key="crono_prazo")

                st.markdown("---")
                st.markdown("**Etapas / Fases de Execução**")
                n_etapas = st.number_input("Número de etapas", min_value=1, max_value=30, value=4, key="crono_n_etapas")
                etapas = []
                for i in range(int(n_etapas)):
                    st.markdown(f"**Etapa {i + 1}**")
                    ec1, ec2, ec3, ec4 = st.columns([3, 1.5, 1.5, 1.5])
                    with ec1:
                        et_desc = st.text_input("Descrição da etapa", key=f"crono_et_desc_{i}")
                    with ec2:
                        et_inicio_mes = st.number_input("Mês início", min_value=1, max_value=120, value=i + 1, key=f"crono_et_ini_{i}")
                    with ec3:
                        et_fim_mes = st.number_input("Mês fim", min_value=1, max_value=120, value=i + 3, key=f"crono_et_fim_{i}")
                    with ec4:
                        et_valor = st.number_input("Valor (R$)", min_value=0.0, format="%.2f", key=f"crono_et_val_{i}")
                    etapas.append({"descricao": et_desc, "mes_inicio": et_inicio_mes, "mes_fim": et_fim_mes, "valor": et_valor})

                crono_obs = st.text_area("Observações", height=60, key="crono_obs_manual")
                submit_crono = st.form_submit_button("📅 Gerar Cronograma", use_container_width=True)

            if submit_crono:
                # Montar documento
                corpo_crono = f"""## CRONOGRAMA FÍSICO-FINANCEIRO

### 1. Dados Gerais
- **Objeto:** {crono_objeto}
- **Nº Contrato/Processo:** {crono_contrato}
- **Empresa Contratada:** {crono_contratada}
- **Valor Global:** {_fmt_brl(crono_valor_global)}
- **Data de Início:** {crono_inicio.strftime('%d/%m/%Y')}
- **Prazo de Execução:** {crono_prazo_meses} meses

### 2. Detalhamento das Etapas

| # | Etapa | Mês Início | Mês Fim | Duração | Valor (R$) | % do Total |
|---|-------|-----------|---------|---------|-----------|-----------|
"""
                soma_etapas = sum(e["valor"] for e in etapas)
                for j, e in enumerate(etapas, 1):
                    dur = max(1, e["mes_fim"] - e["mes_inicio"] + 1)
                    pct = (e["valor"] / crono_valor_global * 100) if crono_valor_global > 0 else 0
                    corpo_crono += f"| {j} | {e['descricao']} | {e['mes_inicio']}° | {e['mes_fim']}° | {dur} meses | {_fmt_brl(e['valor'])} | {pct:.1f}% |\n"

                corpo_crono += f"\n**Total das Etapas:** {_fmt_brl(soma_etapas)}"
                if crono_valor_global > 0:
                    corpo_crono += f" ({soma_etapas / crono_valor_global * 100:.1f}% do valor global)"

                # Cronograma de desembolso mensal
                corpo_crono += "\n\n### 3. Cronograma de Desembolso Mensal\n\n"
                corpo_crono += "| Mês | Etapas Ativas | Desembolso Estimado (R$) | Acumulado (R$) |\n"
                corpo_crono += "|-----|--------------|------------------------|----------------|\n"
                acumulado = 0.0
                for mes in range(1, int(crono_prazo_meses) + 1):
                    desembolso_mes = 0.0
                    ativos = []
                    for e in etapas:
                        if e["mes_inicio"] <= mes <= e["mes_fim"]:
                            dur = max(1, e["mes_fim"] - e["mes_inicio"] + 1)
                            desembolso_mes += e["valor"] / dur
                            ativos.append(e["descricao"][:30])
                    acumulado += desembolso_mes
                    nomes = ", ".join(ativos) if ativos else "—"
                    corpo_crono += f"| {mes}° | {nomes} | {_fmt_brl(desembolso_mes)} | {_fmt_brl(acumulado)} |\n"

                corpo_crono += f"""
### 4. Base Legal
- Lei 14.133/2021, Art. 25, §1° e Art. 46
- IN SEGES/ME nº 58/2022

### 5. Observações
{crono_obs or 'Sem observações.'}

---
Documento gerado pelo sistema O Babilaca (IA) — Ferramenta de apoio.
Confirmar sempre nas fontes oficiais.
"""
                st.markdown(corpo_crono)
                pdf_bytes = gerar_pdf_documento("CRONOGRAMA FÍSICO-FINANCEIRO", corpo_crono)
                docx_bytes = gerar_docx_documento("CRONOGRAMA FÍSICO-FINANCEIRO", corpo_crono)
                _c_pdf, _c_docx = st.columns(2)
                with _c_pdf:
                    st.download_button("⬇️ Baixar Cronograma em PDF", pdf_bytes, "cronograma_fisico_financeiro.pdf", "application/pdf")
                with _c_docx:
                    st.download_button("📝 Baixar Cronograma em Word", docx_bytes, "cronograma_fisico_financeiro.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        else:  # Assistido por IA
            st.markdown(
                "Descreva o projeto/contrato abaixo e a IA gerará um cronograma físico-financeiro "
                "completo com etapas, prazos, valores e curva de desembolso."
            )
            desc_crono = st.text_area(
                "Descreva o objeto, valor estimado, prazo e principais etapas:",
                height=150,
                key="crono_ia_input",
                placeholder=(
                    "Ex: Obra de reforma do prédio administrativo, valor estimado R$ 800.000,00, "
                    "prazo de 12 meses. Inclui demolição, estrutura, instalações elétricas e hidráulicas, "
                    "acabamento e limpeza final."
                ),
            )
            arquivo_crono = st.file_uploader("Ou envie projeto básico/planilha orçamentária (PDF/Excel)", type=["pdf", "xlsx"], key="crono_upload")
            if st.button("🤖 Gerar Cronograma com IA", key="btn_crono_ia"):
                ctx_extra = ""
                if arquivo_crono:
                    if arquivo_crono.name.endswith(".pdf"):
                        ctx_extra = extrair_texto_pdf(arquivo_crono)
                    else:
                        ctx_extra, _ = extrair_dados_excel(arquivo_crono)
                if desc_crono or ctx_extra:
                    with st.spinner("Gerando Cronograma Físico-Financeiro com IA..."):
                        prompt_crono = (
                            "Gere um **Cronograma Físico-Financeiro** completo e detalhado para o seguinte projeto/contrato, "
                            "conforme exigido pela Lei 14.133/2021 (Art. 25, §1° e Art. 46).\n\n"
                            f"**Descrição do projeto:**\n{desc_crono}\n\n"
                        )
                        if ctx_extra:
                            prompt_crono += f"**Dados do arquivo anexado:**\n{ctx_extra[:4000]}\n\n"
                        prompt_crono += (
                            "O documento deve conter OBRIGATORIAMENTE:\n\n"
                            "1. **DADOS GERAIS**: objeto, valor global estimado, prazo total, empresa (quando informada)\n\n"
                            "2. **DETALHAMENTO DAS ETAPAS** em tabela com: Nº, Descrição da Etapa/Fase, "
                            "Mês de Início, Mês de Término, Duração, Valor da Etapa (R$), Percentual do Total (%)\n\n"
                            "3. **CRONOGRAMA DE DESEMBOLSO MENSAL** em tabela com: Mês, Etapas Ativas, "
                            "Desembolso no Mês (R$), Acumulado (R$), Percentual Acumulado (%)\n\n"
                            "4. **CURVA S** descritiva: explique a evolução dos desembolsos ao longo do tempo "
                            "(se há concentração no início, meio ou fim)\n\n"
                            "5. **MARCOS CRÍTICOS**: liste os marcos de medição/pagamento mais importantes\n\n"
                            "6. **CONDIÇÕES DE PAGAMENTO**: vincule os pagamentos à medição das etapas efetivamente executadas\n\n"
                            "7. **BASE LEGAL**: Art. 25 §1° e Art. 46 da Lei 14.133/2021\n\n"
                            "8. **RESSALVAS**: indique que os valores e prazos são estimativos e sujeitos a ajustes\n\n"
                            "Use valores realistas e proporcionais. Formate tabelas em Markdown. "
                            "Todos os valores devem estar em Reais (R$)."
                        )
                        ctx = buscar_base_juridica("cronograma execução pagamento medição prazo contrato obra")
                        system = _build_system_prompt(ctx)
                        resultado = chamar_ia([{"role": "user", "content": prompt_crono}], system)
                    st.markdown(resultado)
                    pdf_bytes = gerar_pdf_documento("CRONOGRAMA FÍSICO-FINANCEIRO", resultado)
                    docx_bytes = gerar_docx_documento("CRONOGRAMA FÍSICO-FINANCEIRO", resultado)
                    _c_pdf, _c_docx = st.columns(2)
                    with _c_pdf:
                        st.download_button("⬇️ Baixar Cronograma em PDF", pdf_bytes, "cronograma_ia.pdf", "application/pdf")
                    with _c_docx:
                        st.download_button("📝 Baixar Cronograma em Word", docx_bytes, "cronograma_ia.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.warning("Descreva o projeto ou envie um arquivo.")

    # ---- MEMÓRIA DE CÁLCULO ----
    elif "Memória de Cálculo" in tipo_doc:
        st.markdown("#### 🧮 Memória de Cálculo")
        st.markdown(
            "Gere uma memória de cálculo detalhada para fundamentar o **valor estimado** da contratação, "
            "conforme **Art. 23 da Lei 14.133/2021** e **IN SEGES/ME nº 65/2021**. "
            "Documento essencial para demonstrar como se chegou ao preço de referência."
        )
        modo_mc = st.radio("Modo de preenchimento", ["Manual", "Assistido por IA"], horizontal=True, key="modo_mc", index=1)

        if modo_mc == "Manual":
            with st.form("form_memoria_calculo"):
                st.markdown("**Dados Gerais**")
                mc_c1, mc_c2 = st.columns(2)
                with mc_c1:
                    mc_objeto = st.text_area("Objeto da contratação", height=60, key="mc_obj")
                    mc_processo = st.text_input("Nº do Processo", key="mc_processo")
                with mc_c2:
                    mc_metodologia = st.selectbox("Metodologia de pesquisa (IN 65/2021)", [
                        "Painel de Preços / ComprasGov",
                        "Contratações similares de outros entes (PNCP)",
                        "Pesquisa direta com fornecedores (mín. 3 cotações)",
                        "Pesquisa em sites especializados / catálogos",
                        "Combinação de fontes (recomendado)",
                    ], key="mc_metodo")
                    mc_criterio = st.selectbox("Critério de formação do preço de referência", [
                        "Média aritmética dos preços válidos",
                        "Mediana dos preços válidos",
                        "Menor preço obtido",
                        "Média saneada (excluídos outliers)",
                    ], key="mc_criterio")

                st.markdown("---")
                st.markdown("**Itens e Cotações**")
                mc_n_itens = st.number_input("Número de itens", min_value=1, max_value=50, value=3, key="mc_n_itens")
                itens_mc = []
                for i in range(int(mc_n_itens)):
                    st.markdown(f"**Item {i + 1}**")
                    ic1, ic2, ic3 = st.columns([3, 1, 1])
                    with ic1:
                        mc_desc = st.text_input("Descrição do item", key=f"mc_desc_{i}")
                    with ic2:
                        mc_und = st.text_input("Unidade", value="UN", key=f"mc_und_{i}")
                    with ic3:
                        mc_qtd = st.number_input("Qtd", min_value=1, value=1, key=f"mc_qtd_{i}")

                    st.markdown("Cotações obtidas:")
                    fc1, fc2, fc3, fc4, fc5 = st.columns(5)
                    with fc1:
                        mc_f1 = st.number_input("Fonte 1 (R$)", min_value=0.0, format="%.2f", key=f"mc_f1_{i}")
                        mc_f1_nome = st.text_input("Nome Fonte 1", key=f"mc_f1n_{i}", placeholder="Ex: Ata PE 10/2025")
                    with fc2:
                        mc_f2 = st.number_input("Fonte 2 (R$)", min_value=0.0, format="%.2f", key=f"mc_f2_{i}")
                        mc_f2_nome = st.text_input("Nome Fonte 2", key=f"mc_f2n_{i}", placeholder="Ex: Cotação Empresa X")
                    with fc3:
                        mc_f3 = st.number_input("Fonte 3 (R$)", min_value=0.0, format="%.2f", key=f"mc_f3_{i}")
                        mc_f3_nome = st.text_input("Nome Fonte 3", key=f"mc_f3n_{i}", placeholder="Ex: Painel de Preços")
                    with fc4:
                        mc_f4 = st.number_input("Fonte 4 (R$)", min_value=0.0, format="%.2f", key=f"mc_f4_{i}")
                        mc_f4_nome = st.text_input("Nome Fonte 4", key=f"mc_f4n_{i}", placeholder="(Opcional)")
                    with fc5:
                        mc_f5 = st.number_input("Fonte 5 (R$)", min_value=0.0, format="%.2f", key=f"mc_f5_{i}")
                        mc_f5_nome = st.text_input("Nome Fonte 5", key=f"mc_f5n_{i}", placeholder="(Opcional)")

                    itens_mc.append({
                        "descricao": mc_desc, "unidade": mc_und, "quantidade": mc_qtd,
                        "fontes": [
                            {"nome": mc_f1_nome, "valor": mc_f1},
                            {"nome": mc_f2_nome, "valor": mc_f2},
                            {"nome": mc_f3_nome, "valor": mc_f3},
                            {"nome": mc_f4_nome, "valor": mc_f4},
                            {"nome": mc_f5_nome, "valor": mc_f5},
                        ],
                    })

                mc_obs = st.text_area("Observações e justificativas adicionais", height=80, key="mc_obs")
                submit_mc = st.form_submit_button("🧮 Gerar Memória de Cálculo", use_container_width=True)

            if submit_mc:
                import statistics
                corpo_mc = f"""## MEMÓRIA DE CÁLCULO — ESTIMATIVA DE PREÇOS

### 1. Dados Gerais
- **Objeto:** {mc_objeto}
- **Nº Processo:** {mc_processo}
- **Metodologia:** {mc_metodologia}
- **Critério de formação do preço:** {mc_criterio}

### 2. Fontes de Pesquisa Utilizadas
Conforme Art. 23, §1° da Lei 14.133/2021 e IN SEGES/ME nº 65/2021, a pesquisa de preços
utilizou as seguintes fontes, priorizando a ordem estabelecida no Art. 5° da IN 65/2021:

"""
                # Coletar todas as fontes únicas dos itens
                fontes_unicas = set()
                for item in itens_mc:
                    for f in item["fontes"]:
                        if f["nome"] and f["valor"] > 0:
                            fontes_unicas.add(f["nome"])
                for idx_f, fn in enumerate(sorted(fontes_unicas), 1):
                    corpo_mc += f"- **Fonte {idx_f}:** {fn}\n"

                corpo_mc += "\n### 3. Detalhamento por Item\n\n"

                valor_total_global = 0.0
                for j, item in enumerate(itens_mc, 1):
                    fontes_validas = [f for f in item["fontes"] if f["valor"] > 0]
                    precos = [f["valor"] for f in fontes_validas]

                    corpo_mc += f"---\n#### Item {j}: {item['descricao']}\n"
                    corpo_mc += f"- **Unidade:** {item['unidade']} | **Quantidade:** {item['quantidade']}\n\n"

                    # Tabela de cotações
                    corpo_mc += "| Fonte | Valor Unitário (R$) |\n"
                    corpo_mc += "|-------|--------------------|\n"
                    for f in fontes_validas:
                        corpo_mc += f"| {f['nome']} | {_fmt_brl(f['valor'])} |\n"

                    if precos:
                        media = sum(precos) / len(precos)
                        mediana = statistics.median(precos)
                        menor = min(precos)
                        maior = max(precos)
                        coef_var = (statistics.stdev(precos) / media * 100) if len(precos) > 1 else 0

                        # Determinar preço de referência conforme critério
                        if "Mediana" in mc_criterio:
                            preco_ref = mediana
                        elif "Menor" in mc_criterio:
                            preco_ref = menor
                        elif "saneada" in mc_criterio.lower():
                            # Excluir valores fora de 1 desvio padrão
                            if len(precos) > 2:
                                dp = statistics.stdev(precos)
                                saneados = [p for p in precos if abs(p - media) <= dp]
                                preco_ref = sum(saneados) / len(saneados) if saneados else media
                            else:
                                preco_ref = media
                        else:
                            preco_ref = media

                        total_item = preco_ref * item["quantidade"]
                        valor_total_global += total_item

                        corpo_mc += f"\n**Análise Estatística:**\n"
                        corpo_mc += f"- Menor preço: {_fmt_brl(menor)}\n"
                        corpo_mc += f"- Maior preço: {_fmt_brl(maior)}\n"
                        corpo_mc += f"- Média aritmética: {_fmt_brl(media)}\n"
                        corpo_mc += f"- Mediana: {_fmt_brl(mediana)}\n"
                        if len(precos) > 1:
                            corpo_mc += f"- Coeficiente de variação: {coef_var:.1f}%"
                            if coef_var > 25:
                                corpo_mc += " ⚠️ *Dispersão alta — recomenda-se justificativa*"
                            corpo_mc += "\n"
                        corpo_mc += f"\n**▶ Preço de Referência Unitário ({mc_criterio}):** {_fmt_brl(preco_ref)}\n"
                        corpo_mc += f"**▶ Valor Total do Item ({item['quantidade']} x {_fmt_brl(preco_ref)}):** {_fmt_brl(total_item)}\n\n"
                    else:
                        corpo_mc += "\n⚠️ *Nenhuma cotação informada para este item.*\n\n"

                corpo_mc += f"""---
### 4. Resumo — Valor Total Estimado

| Critério | Valor |
|----------|-------|
| **Valor Total Estimado Global** | **{_fmt_brl(valor_total_global)}** |
| Metodologia aplicada | {mc_metodologia} |
| Critério de formação | {mc_criterio} |
| Nº de fontes utilizadas | {len(fontes_unicas)} |

### 5. Conformidade Legal
- **Art. 23, Lei 14.133/2021** — Valor estimado compatível com mercado
- **IN SEGES/ME nº 65/2021, Arts. 5° a 7°** — Parâmetros e ordem de fontes
- **Acórdão 1.793/2011 — TCU Plenário** — Pesquisa ampla e diversificada

### 6. Observações
{mc_obs or 'Sem observações adicionais.'}

---
Documento gerado pelo sistema O Babilaca (IA) — Ferramenta de apoio.
Confirmar sempre nas fontes oficiais.
"""
                st.markdown(corpo_mc)
                pdf_bytes = gerar_pdf_documento("MEMÓRIA DE CÁLCULO", corpo_mc)
                docx_bytes = gerar_docx_documento("MEMÓRIA DE CÁLCULO", corpo_mc)
                _c_pdf, _c_docx = st.columns(2)
                with _c_pdf:
                    st.download_button("⬇️ Baixar Memória de Cálculo em PDF", pdf_bytes, "memoria_calculo.pdf", "application/pdf")
                with _c_docx:
                    st.download_button("📝 Baixar Memória de Cálculo em Word", docx_bytes, "memoria_calculo.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        else:  # Assistido por IA
            st.markdown(
                "Descreva os itens a contratar, fontes de pesquisa e valores obtidos. "
                "A IA gerará uma memória de cálculo completa com análise estatística, "
                "justificativa e fundamentação legal."
            )
            desc_mc = st.text_area(
                "Descreva os itens, quantidades, valores pesquisados e fontes utilizadas:",
                height=180,
                key="mc_ia_input",
                placeholder=(
                    "Ex: Preciso de memória de cálculo para aquisição de 50 notebooks Dell Latitude 5540.\n"
                    "Fontes pesquisadas:\n"
                    "- Ata PE 15/2025 UASG 160XXX: R$ 5.200,00/un\n"
                    "- Painel de Preços (ComprasGov): R$ 5.450,00/un\n"
                    "- Cotação Fornecedor ABC: R$ 5.100,00/un\n"
                    "- Site Magazine Luiza: R$ 5.800,00/un"
                ),
            )
            arquivo_mc = st.file_uploader("Ou envie planilha de cotações / pesquisa de preços (PDF/Excel)", type=["pdf", "xlsx"], key="mc_upload")
            if st.button("🤖 Gerar Memória de Cálculo com IA", key="btn_mc_ia"):
                ctx_extra = ""
                if arquivo_mc:
                    if arquivo_mc.name.endswith(".pdf"):
                        ctx_extra = extrair_texto_pdf(arquivo_mc)
                    else:
                        ctx_extra, _ = extrair_dados_excel(arquivo_mc)
                if desc_mc or ctx_extra:
                    with st.spinner("Gerando Memória de Cálculo com IA..."):
                        prompt_mc = (
                            "Gere uma **Memória de Cálculo** completa e detalhada para fundamentar o valor estimado "
                            "da contratação abaixo, conforme Art. 23 da Lei 14.133/2021 e IN SEGES/ME nº 65/2021.\n\n"
                            f"**Descrição / Dados fornecidos:**\n{desc_mc}\n\n"
                        )
                        if ctx_extra:
                            prompt_mc += f"**Dados do arquivo anexado:**\n{ctx_extra[:4000]}\n\n"
                        prompt_mc += (
                            "O documento deve conter OBRIGATORIAMENTE:\n\n"
                            "1. **DADOS GERAIS**: objeto, processo (se informado), metodologia de pesquisa\n\n"
                            "2. **FONTES DE PESQUISA**: listar todas as fontes utilizadas, seguindo a ordem "
                            "de prioridade da IN 65/2021 Art. 5°: I-Painel de Preços, II-Contratações similares, "
                            "III-Sites especializados, IV-Pesquisa direta com fornecedores\n\n"
                            "3. **DETALHAMENTO POR ITEM** com tabela de cotações contendo: Fonte, CNPJ/Identificação "
                            "(quando disponível), Valor Unitário, Data da Cotação\n\n"
                            "4. **ANÁLISE ESTATÍSTICA POR ITEM**: média aritmética, mediana, menor preço, maior preço, "
                            "desvio padrão, coeficiente de variação. Se CV > 25%, alertar sobre dispersão alta\n\n"
                            "5. **TRATAMENTO DE OUTLIERS**: se houver preços com dispersão excessiva (>25%), "
                            "aplicar média saneada excluindo valores fora de 1 desvio padrão, justificando a exclusão\n\n"
                            "6. **PREÇO DE REFERÊNCIA**: indicar o critério adotado (média, mediana ou menor), "
                            "calcular o valor unitário e o valor total por item e global\n\n"
                            "7. **QUADRO RESUMO** com tabela: Item, Qtd, Unidade, Preço Ref. Unitário, Valor Total\n\n"
                            "8. **CONFORMIDADE LEGAL**: citar Art. 23 Lei 14.133, IN 65/2021, e jurisprudência "
                            "relevante (Acórdão 1.793/2011-TCU caso esteja no contexto)\n\n"
                            "9. **CONCLUSÃO**: valor total estimado e declaração de compatibilidade com mercado\n\n"
                            "Use tabelas Markdown, valores em Reais (R$) formatados, e cálculos demonstrados passo a passo."
                        )
                        ctx = buscar_base_juridica("pesquisa preços estimativa orçamento cotação memória cálculo IN 65")
                        system = _build_system_prompt(ctx)
                        resultado = chamar_ia([{"role": "user", "content": prompt_mc}], system)
                    st.markdown(resultado)
                    pdf_bytes = gerar_pdf_documento("MEMÓRIA DE CÁLCULO", resultado)
                    docx_bytes = gerar_docx_documento("MEMÓRIA DE CÁLCULO", resultado)
                    _c_pdf, _c_docx = st.columns(2)
                    with _c_pdf:
                        st.download_button("⬇️ Baixar Memória de Cálculo em PDF", pdf_bytes, "memoria_calculo_ia.pdf", "application/pdf")
                    with _c_docx:
                        st.download_button("📝 Baixar Memória de Cálculo em Word", docx_bytes, "memoria_calculo_ia.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.warning("Descreva os itens/valores ou envie um arquivo.")

    # ---- ATA DE REGISTRO DE PREÇOS ----
    elif "Ata de Registro" in tipo_doc:
        st.markdown("#### 📜 Modelo de Ata de Registro de Preços")
        st.markdown(
            "Gere um modelo completo de Ata de Registro de Preços conforme a **Lei 14.133/2021 (Arts. 82 a 86)** "
            "e o **Decreto nº 11.462/2023**. Inclui cláusulas de revisão, cancelamento, adesão (carona), "
            "obrigações e penalidades."
        )
        modo_ata = st.radio("Modo de preenchimento", ["Manual", "Assistido por IA"], horizontal=True, key="modo_ata", index=1)

        if modo_ata == "Manual":
            with st.form("form_ata_rp"):
                st.markdown("**Dados do Órgão Gerenciador**")
                ca1, ca2 = st.columns(2)
                with ca1:
                    ata_orgao = st.text_input("Órgão Gerenciador", key="ata_orgao")
                    ata_cnpj_orgao = st.text_input("CNPJ do Órgão", key="ata_cnpj_orgao")
                    ata_endereco_orgao = st.text_input("Endereço do Órgão", key="ata_end_orgao")
                with ca2:
                    ata_representante = st.text_input("Representante Legal", key="ata_repr")
                    ata_cargo = st.text_input("Cargo do Representante", key="ata_cargo")
                    ata_processo = st.text_input("Nº Processo Administrativo", key="ata_processo")

                st.markdown("---")
                st.markdown("**Dados da Licitação**")
                cb1, cb2 = st.columns(2)
                with cb1:
                    ata_pregao = st.text_input("Nº do Pregão Eletrônico", key="ata_pregao")
                    ata_objeto = st.text_area("Objeto", height=60, key="ata_objeto")
                with cb2:
                    ata_vigencia = st.text_input("Vigência da Ata", value="12 (doze) meses", key="ata_vigencia")
                    ata_data_assinatura = st.date_input("Data de Assinatura", key="ata_data_ass")

                st.markdown("---")
                st.markdown("**Dados do Fornecedor**")
                cc1, cc2 = st.columns(2)
                with cc1:
                    ata_fornecedor = st.text_input("Razão Social do Fornecedor", key="ata_forn")
                    ata_cnpj_forn = st.text_input("CNPJ do Fornecedor", key="ata_cnpj_forn")
                    ata_end_forn = st.text_input("Endereço do Fornecedor", key="ata_end_forn")
                with cc2:
                    ata_repr_forn = st.text_input("Representante do Fornecedor", key="ata_repr_forn")
                    ata_contato_forn = st.text_input("Telefone/E-mail do Fornecedor", key="ata_contato_forn")

                st.markdown("---")
                st.markdown("**Itens Registrados**")
                n_itens_ata = st.number_input("Número de itens", min_value=1, max_value=50, value=3, key="n_itens_ata")
                itens_ata = []
                for i in range(int(n_itens_ata)):
                    st.markdown(f"**Item {i + 1}**")
                    ia1, ia2, ia3, ia4 = st.columns([4, 1, 1, 1.5])
                    with ia1:
                        ata_desc = st.text_input("Descrição", key=f"ata_desc_{i}")
                    with ia2:
                        ata_und = st.text_input("Unidade", value="UN", key=f"ata_und_{i}")
                    with ia3:
                        ata_qtd = st.number_input("Qtd", min_value=1, value=1, key=f"ata_qtd_{i}")
                    with ia4:
                        ata_val = st.number_input("Valor Unit. (R$)", min_value=0.0, format="%.2f", key=f"ata_val_{i}")
                    itens_ata.append({"descricao": ata_desc, "unidade": ata_und, "quantidade": ata_qtd, "valor": ata_val})

                st.markdown("---")
                st.markdown("**Condições**")
                cd1, cd2 = st.columns(2)
                with cd1:
                    ata_prazo_entrega = st.text_input("Prazo de entrega", value="30 (trinta) dias", key="ata_prazo_ent")
                    ata_local_entrega = st.text_input("Local de entrega", key="ata_local_ent")
                with cd2:
                    ata_prazo_pgto = st.text_input("Prazo de pagamento", value="30 (trinta) dias", key="ata_prazo_pgto")
                ata_obs = st.text_area("Observações / Disposições Gerais", height=80, key="ata_obs")
                submit_ata = st.form_submit_button("📜 Gerar Ata de Registro de Preços", use_container_width=True)

            if submit_ata:
                tabela_itens_str = ""
                valor_total_ata = 0.0
                for j, it in enumerate(itens_ata, 1):
                    vt = it["valor"] * it["quantidade"]
                    valor_total_ata += vt
                    tabela_itens_str += f"| {j} | {it['descricao']} | {it['unidade']} | {it['quantidade']} | {_fmt_brl(it['valor'])} | {_fmt_brl(vt)} |\n"

                corpo_ata = TEMPLATE_ATA_RP.format(
                    processo=ata_processo, pregao=ata_pregao, orgao=ata_orgao,
                    cnpj_orgao=ata_cnpj_orgao, vigencia=ata_vigencia,
                    data_assinatura=ata_data_assinatura.strftime("%d/%m/%Y"),
                    endereco_orgao=ata_endereco_orgao, cargo_representante=ata_cargo,
                    representante=ata_representante, objeto=ata_objeto,
                    fornecedor=ata_fornecedor, cnpj_fornecedor=ata_cnpj_forn,
                    endereco_fornecedor=ata_end_forn, representante_fornecedor=ata_repr_forn,
                    contato_fornecedor=ata_contato_forn, tabela_itens=tabela_itens_str,
                    valor_total=_fmt_brl(valor_total_ata), prazo_entrega=ata_prazo_entrega,
                    local_entrega=ata_local_entrega, prazo_pagamento=ata_prazo_pgto,
                    observacoes=ata_obs or "Sem observações adicionais.",
                )
                st.markdown(corpo_ata)
                pdf_bytes = gerar_pdf_documento("ATA DE REGISTRO DE PREÇOS", corpo_ata)
                docx_bytes = gerar_docx_documento("ATA DE REGISTRO DE PREÇOS", corpo_ata)
                _c_pdf, _c_docx = st.columns(2)
                with _c_pdf:
                    st.download_button("⬇️ Baixar Ata em PDF", pdf_bytes, "ata_registro_precos.pdf", "application/pdf")
                with _c_docx:
                    st.download_button("📝 Baixar Ata em Word", docx_bytes, "ata_registro_precos.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            desc_ata = st.text_area(
                "Descreva os dados da licitação, órgão, fornecedor, itens e condições. O Babilaca gerará a Ata completa:",
                height=180,
                key="ata_ia_input",
                placeholder=(
                    "Ex: Pregão Eletrônico nº 10/2026, UASG 787000 (Centro de Obtenção da Marinha).\n"
                    "Fornecedor vencedor: Empresa ABC Ltda, CNPJ 12.345.678/0001-90.\n"
                    "Itens: 100 un resma de papel A4 a R$ 25,00/un; 50 un toner HP a R$ 180,00/un.\n"
                    "Vigência: 12 meses. Entrega em 30 dias no almoxarifado central."
                ),
            )
            arquivo_ata = st.file_uploader("Ou envie edital/resultado de licitação (PDF/Excel)", type=["pdf", "xlsx"], key="ata_upload")
            if st.button("🤖 Gerar Ata com IA", key="btn_ata_ia"):
                ctx_extra = ""
                if arquivo_ata:
                    if arquivo_ata.name.endswith(".pdf"):
                        ctx_extra = extrair_texto_pdf(arquivo_ata)
                    else:
                        ctx_extra, _ = extrair_dados_excel(arquivo_ata)
                if desc_ata or ctx_extra:
                    with st.spinner("Gerando Ata de Registro de Preços com IA..."):
                        prompt_ata = (
                            "Gere um **Modelo de Ata de Registro de Preços** completo e detalhado, "
                            "conforme a Lei 14.133/2021 (Arts. 82 a 86) e o Decreto 11.462/2023.\n\n"
                            f"**Dados fornecidos:**\n{desc_ata}\n\n"
                        )
                        if ctx_extra:
                            prompt_ata += f"**Dados do arquivo anexado:**\n{ctx_extra[:4000]}\n\n"
                        prompt_ata += (
                            "O documento deve conter OBRIGATORIAMENTE:\n\n"
                            "1. **PREÂMBULO** com identificação do Órgão Gerenciador, licitação e fornecedor\n"
                            "2. **OBJETO** detalhado\n"
                            "3. **TABELA DE ITENS** com descrição, unidade, quantidade, valor unitário e total\n"
                            "4. **VALIDADE DA ATA** — prazo e condições de prorrogação (Art. 84)\n"
                            "5. **REVISÃO E CANCELAMENTO** de preços (Art. 82, §5°)\n"
                            "6. **CONDIÇÕES DE ENTREGA** — prazo, local, recebimento provisório/definitivo\n"
                            "7. **PAGAMENTO** — prazo, documentação, retenções\n"
                            "8. **OBRIGAÇÕES DO FORNECEDOR** — habilitação, encargos, substituição\n"
                            "9. **OBRIGAÇÕES DO ÓRGÃO** — gerenciamento, renegociação, sanções\n"
                            "10. **ADESÃO (CARONA)** — limites de 50% por órgão e dobro total (Art. 86)\n"
                            "11. **PENALIDADES** — Arts. 155 a 163\n"
                            "12. **FUNDAMENTAÇÃO LEGAL** — Lei 14.133/2021, Decreto 11.462/2023\n\n"
                            "Use tabelas Markdown, valores em Reais (R$) e linguagem jurídica formal."
                        )
                        ctx = buscar_base_juridica("ata registro preços pregão licitação fornecedor")
                        system = _build_system_prompt(ctx)
                        resultado = chamar_ia([{"role": "user", "content": prompt_ata}], system)
                    st.markdown(resultado)
                    pdf_bytes = gerar_pdf_documento("ATA DE REGISTRO DE PREÇOS", resultado)
                    docx_bytes = gerar_docx_documento("ATA DE REGISTRO DE PREÇOS", resultado)
                    _c_pdf, _c_docx = st.columns(2)
                    with _c_pdf:
                        st.download_button("⬇️ Baixar Ata em PDF", pdf_bytes, "ata_rp_ia.pdf", "application/pdf")
                    with _c_docx:
                        st.download_button("📝 Baixar Ata em Word", docx_bytes, "ata_rp_ia.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.warning("Descreva os dados da licitação ou envie um arquivo.")

    # ---- MINUTA DE CONTRATO ----
    elif "Minuta de Contrato" in tipo_doc:
        st.markdown("#### 📑 Minuta de Contrato de Serviço / Material")
        st.markdown(
            "Gere uma minuta de contrato completa com todas as cláusulas essenciais, incluindo "
            "**garantia contratual**, **garantia dos produtos/serviços**, penalidades, fiscalização "
            "e recebimento, conforme a **Lei 14.133/2021**."
        )
        modo_contrato = st.radio("Modo de preenchimento", ["Manual", "Assistido por IA"], horizontal=True, key="modo_contrato", index=1)

        if modo_contrato == "Manual":
            with st.form("form_minuta_contrato"):
                st.markdown("**Dados do Contratante**")
                mc1, mc2 = st.columns(2)
                with mc1:
                    ct_orgao = st.text_input("Órgão/Entidade", key="ct_orgao")
                    ct_cnpj_orgao = st.text_input("CNPJ do Órgão", key="ct_cnpj_orgao")
                    ct_endereco_orgao = st.text_input("Endereço do Órgão", key="ct_end_orgao")
                with mc2:
                    ct_representante = st.text_input("Representante Legal", key="ct_repr")
                    ct_cargo = st.text_input("Cargo", key="ct_cargo")
                    ct_processo = st.text_input("Nº Processo Administrativo", key="ct_processo")

                st.markdown("---")
                st.markdown("**Dados da Licitação e Contrato**")
                md1, md2 = st.columns(2)
                with md1:
                    ct_licitacao = st.text_input("Nº da Licitação", key="ct_licitacao")
                    ct_modalidade = st.selectbox("Modalidade", [
                        "Pregão Eletrônico", "Concorrência", "Dispensa de Licitação",
                        "Inexigibilidade", "Adesão a Ata de Registro de Preços"
                    ], key="ct_modalidade")
                    ct_tipo = st.selectbox("Tipo de Contrato", [
                        "FORNECIMENTO DE MATERIAL", "PRESTAÇÃO DE SERVIÇO",
                        "FORNECIMENTO DE MATERIAL E PRESTAÇÃO DE SERVIÇO"
                    ], key="ct_tipo")
                with md2:
                    ct_ano = st.text_input("Ano", value=str(datetime.now().year), key="ct_ano")
                    ct_vigencia = st.text_input("Vigência", value="12 (doze) meses", key="ct_vigencia")
                    ct_objeto = st.text_area("Objeto da Contratação", height=60, key="ct_objeto")

                st.markdown("---")
                st.markdown("**Dados da Contratada**")
                me1, me2 = st.columns(2)
                with me1:
                    ct_empresa = st.text_input("Razão Social da Empresa", key="ct_empresa")
                    ct_cnpj_empresa = st.text_input("CNPJ da Empresa", key="ct_cnpj_emp")
                    ct_end_empresa = st.text_input("Endereço da Empresa", key="ct_end_emp")
                with me2:
                    ct_repr_empresa = st.text_input("Representante da Empresa", key="ct_repr_emp")

                st.markdown("---")
                st.markdown("**Itens / Serviços**")
                n_itens_ct = st.number_input("Número de itens", min_value=1, max_value=50, value=3, key="n_itens_ct")
                itens_ct = []
                for i in range(int(n_itens_ct)):
                    st.markdown(f"**Item {i + 1}**")
                    ic1, ic2, ic3, ic4 = st.columns([4, 1, 1, 1.5])
                    with ic1:
                        ct_desc = st.text_input("Descrição", key=f"ct_desc_{i}")
                    with ic2:
                        ct_und = st.text_input("Unidade", value="UN", key=f"ct_und_{i}")
                    with ic3:
                        ct_qtd = st.number_input("Qtd", min_value=1, value=1, key=f"ct_qtd_{i}")
                    with ic4:
                        ct_val = st.number_input("Valor Unit. (R$)", min_value=0.0, format="%.2f", key=f"ct_val_{i}")
                    itens_ct.append({"descricao": ct_desc, "unidade": ct_und, "quantidade": ct_qtd, "valor": ct_val})

                st.markdown("---")
                st.markdown("**Garantia e Condições**")
                gf1, gf2 = st.columns(2)
                with gf1:
                    ct_garantia_contratual = st.selectbox("Garantia contratual (% do valor)", [
                        "5%", "10%", "Até 30% (obras de grande vulto)"
                    ], key="ct_gar_contratual")
                    ct_garantia_produto = st.text_input("Garantia dos produtos/serviços", value="12 (doze) meses", key="ct_gar_prod")
                    ct_prazo_reparo = st.text_input("Prazo para reparo/substituição em garantia", value="10 (dez) dias úteis", key="ct_prazo_reparo")
                with gf2:
                    ct_prazo_execucao = st.text_input("Prazo de execução/entrega", value="30 (trinta) dias", key="ct_prazo_exec")
                    ct_prazo_pgto = st.text_input("Prazo de pagamento", value="30 (trinta) dias", key="ct_prazo_pgto")

                st.markdown("---")
                st.markdown("**Penalidades**")
                gp1, gp2, gp3 = st.columns(3)
                with gp1:
                    ct_multa_mora = st.text_input("Multa moratória (por dia)", value="0,5%", key="ct_multa_mora")
                with gp2:
                    ct_limite_multa = st.text_input("Limite da multa moratória", value="10%", key="ct_limite_multa")
                with gp3:
                    ct_multa_comp = st.text_input("Multa compensatória", value="20%", key="ct_multa_comp")

                st.markdown("---")
                st.markdown("**Recebimento**")
                gr1, gr2 = st.columns(2)
                with gr1:
                    ct_receb_prov = st.text_input("Prazo receb. provisório", value="5 (cinco) dias úteis", key="ct_rec_prov")
                with gr2:
                    ct_receb_def = st.text_input("Prazo receb. definitivo", value="10 (dez) dias úteis", key="ct_rec_def")

                ct_obs = st.text_area("Observações / Disposições Gerais", height=80, key="ct_obs")
                submit_ct = st.form_submit_button("📑 Gerar Minuta de Contrato", use_container_width=True)

            if submit_ct:
                tabela_itens_ct = ""
                valor_total_ct = 0.0
                for j, it in enumerate(itens_ct, 1):
                    vt = it["valor"] * it["quantidade"]
                    valor_total_ct += vt
                    tabela_itens_ct += f"| {j} | {it['descricao']} | {it['unidade']} | {it['quantidade']} | {_fmt_brl(it['valor'])} | {_fmt_brl(vt)} |\n"

                corpo_ct = TEMPLATE_MINUTA_CONTRATO.format(
                    tipo_contrato=ct_tipo, processo=ct_processo, licitacao=ct_licitacao,
                    modalidade=ct_modalidade, ano=ct_ano, vigencia=ct_vigencia,
                    orgao=ct_orgao, endereco_orgao=ct_endereco_orgao, cnpj_orgao=ct_cnpj_orgao,
                    cargo_representante=ct_cargo, representante=ct_representante,
                    empresa=ct_empresa, cnpj_empresa=ct_cnpj_empresa,
                    endereco_empresa=ct_end_empresa, representante_empresa=ct_repr_empresa,
                    objeto=ct_objeto, tabela_itens=tabela_itens_ct,
                    valor_total=_fmt_brl(valor_total_ct),
                    valor_extenso="(valor por extenso)",
                    prazo_pagamento=ct_prazo_pgto, prazo_execucao=ct_prazo_execucao,
                    percentual_garantia_contratual=ct_garantia_contratual,
                    garantia_produto=ct_garantia_produto, prazo_reparo=ct_prazo_reparo,
                    multa_moratoria=ct_multa_mora, limite_multa=ct_limite_multa,
                    multa_compensatoria=ct_multa_comp,
                    prazo_recebimento_provisorio=ct_receb_prov,
                    prazo_recebimento_definitivo=ct_receb_def,
                    observacoes=ct_obs or "Sem observações adicionais.",
                )
                st.markdown(corpo_ct)
                pdf_bytes = gerar_pdf_documento(f"MINUTA DE CONTRATO DE {ct_tipo}", corpo_ct)
                docx_bytes = gerar_docx_documento(f"MINUTA DE CONTRATO DE {ct_tipo}", corpo_ct)
                _c_pdf, _c_docx = st.columns(2)
                with _c_pdf:
                    st.download_button("⬇️ Baixar Minuta em PDF", pdf_bytes, "minuta_contrato.pdf", "application/pdf")
                with _c_docx:
                    st.download_button("📝 Baixar Minuta em Word", docx_bytes, "minuta_contrato.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            desc_ct = st.text_area(
                "Descreva o objeto, tipo (serviço/material), condições e o Babilaca gerará a Minuta completa:",
                height=180,
                key="ct_ia_input",
                placeholder=(
                    "Ex: Contrato de fornecimento de 200 cadeiras ergonômicas para o prédio administrativo.\n"
                    "Empresa: Móveis XYZ Ltda. Valor: R$ 150.000,00. Prazo de entrega: 45 dias.\n"
                    "Garantia do produto: 5 anos. Garantia contratual: 5% em seguro-garantia.\n"
                    "Incluir cláusula de assistência técnica e substituição em caso de defeito."
                ),
            )
            arquivo_ct = st.file_uploader("Ou envie Termo de Referência / proposta (PDF/Excel)", type=["pdf", "xlsx"], key="ct_upload")
            if st.button("🤖 Gerar Minuta com IA", key="btn_ct_ia"):
                ctx_extra = ""
                if arquivo_ct:
                    if arquivo_ct.name.endswith(".pdf"):
                        ctx_extra = extrair_texto_pdf(arquivo_ct)
                    else:
                        ctx_extra, _ = extrair_dados_excel(arquivo_ct)
                if desc_ct or ctx_extra:
                    with st.spinner("Gerando Minuta de Contrato com IA..."):
                        prompt_ct = (
                            "Gere uma **Minuta de Contrato** completa e detalhada para a contratação abaixo, "
                            "conforme a Lei 14.133/2021.\n\n"
                            f"**Dados fornecidos:**\n{desc_ct}\n\n"
                        )
                        if ctx_extra:
                            prompt_ct += f"**Dados do arquivo anexado:**\n{ctx_extra[:4000]}\n\n"
                        prompt_ct += (
                            "O documento deve conter OBRIGATORIAMENTE todas as cláusulas essenciais:\n\n"
                            "1. **PREÂMBULO** — identificação das partes (CONTRATANTE e CONTRATADA)\n"
                            "2. **OBJETO** — descrição detalhada\n"
                            "3. **ITENS/ESPECIFICAÇÕES** — tabela com descrição, unidade, qtd, valor unitário e total\n"
                            "4. **VALOR E PAGAMENTO** — valor total, prazo de pagamento, condições, retenções fiscais\n"
                            "5. **VIGÊNCIA E EXECUÇÃO** — prazos, possibilidade de prorrogação (Art. 107)\n"
                            "6. **GARANTIA CONTRATUAL** (CLÁUSULA ESSENCIAL!) — percentual (5% a 10%), "
                            "modalidades (caução, seguro-garantia, fiança bancária), prazo de apresentação, "
                            "condições de liberação (Art. 96 da Lei 14.133/2021)\n"
                            "7. **GARANTIA DOS PRODUTOS/SERVIÇOS** (CLÁUSULA ESSENCIAL!) — prazo mínimo de garantia, "
                            "obrigações de substituição/reparo, prazo para atendimento de chamados, "
                            "cobertura (defeitos de fabricação, vícios, não conformidade), "
                            "exclusões (mau uso, caso fortuito), suspensão do prazo durante reparo, "
                            "assistência técnica, reexecução de serviços defeituosos\n"
                            "8. **OBRIGAÇÕES DA CONTRATADA** — habilitação, encargos, preposto, EPI, segurança\n"
                            "9. **OBRIGAÇÕES DO CONTRATANTE** — fiscalização, pagamento, comunicação\n"
                            "10. **FISCALIZAÇÃO** — fiscal e gestor do contrato (Art. 117)\n"
                            "11. **RECEBIMENTO** — provisório e definitivo, prazos, conferência\n"
                            "12. **ALTERAÇÃO CONTRATUAL** — limites de 25%/50% (Arts. 124-136)\n"
                            "13. **PENALIDADES** — multa moratória (% por dia), multa compensatória, "
                            "impedimento de licitar, inidoneidade (Arts. 155-163)\n"
                            "14. **RESCISÃO** — hipóteses (Arts. 137-139)\n"
                            "15. **FUNDAMENTAÇÃO LEGAL**\n"
                            "16. **FORO**\n"
                            "17. **ASSINATURAS** — espaço para CONTRATANTE, CONTRATADA e testemunhas\n\n"
                            "ATENÇÃO ESPECIAL às cláusulas de GARANTIA: sejam extremamente detalhistas sobre "
                            "prazos, coberturas, exclusões, procedimentos de acionamento e penalidades por descumprimento.\n\n"
                            "Use linguagem jurídica formal, tabelas Markdown e valores em Reais (R$)."
                        )
                        ctx = buscar_base_juridica("contrato garantia fiscalização recebimento penalidade pagamento")
                        system = _build_system_prompt(ctx)
                        resultado = chamar_ia([{"role": "user", "content": prompt_ct}], system)
                    st.markdown(resultado)
                    pdf_bytes = gerar_pdf_documento("MINUTA DE CONTRATO", resultado)
                    docx_bytes = gerar_docx_documento("MINUTA DE CONTRATO", resultado)
                    _c_pdf, _c_docx = st.columns(2)
                    with _c_pdf:
                        st.download_button("⬇️ Baixar Minuta em PDF", pdf_bytes, "minuta_contrato_ia.pdf", "application/pdf")
                    with _c_docx:
                        st.download_button("📝 Baixar Minuta em Word", docx_bytes, "minuta_contrato_ia.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.warning("Descreva a contratação ou envie um arquivo.")

    # ---- Classificação de Risco Jurídico ----
    elif "Risco" in tipo_doc:
        st.markdown("#### ⚖️ Classificação de Risco Jurídico")
        st.markdown(
            "Descreva o cenário de contratação e a IA irá gerar uma **classificação de risco** "
            "com **mapa de riscos** (probabilidade × impacto) e **recomendações de mitigação**."
        )
        cenario = st.text_area(
            "Descreva o cenário de contratação para análise de risco:",
            height=150,
            key="cenario_risco",
            placeholder="Ex: Contratação direta de software de gestão por dispensa de licitação, valor estimado R$ 70.000...",
        )
        if st.button("⚖️ Analisar Risco Jurídico", key="btn_risco", use_container_width=True):
            if cenario:
                with st.spinner("Analisando risco jurídico e gerando mapa de riscos..."):
                    resultado = classificar_risco_juridico(cenario)
                st.markdown("---")
                st.markdown(resultado)
                # Legenda visual da matriz de risco
                st.markdown("---")
                st.markdown("##### 🗺️ Legenda — Matriz de Risco")
                legenda_html = """
                <div style="margin-top:0.5rem;">
                <table style="border-collapse:collapse;width:100%;max-width:600px;margin:0 auto;text-align:center;font-size:13px;color:#e0e0e0;">
                <thead>
                <tr style="background:#001a4d;">
                    <th style="border:1px solid #d4af37;padding:8px 10px;color:#d4af37;">Probabilidade \\ Impacto</th>
                    <th style="border:1px solid #d4af37;padding:8px 10px;color:#d4af37;">Baixo</th>
                    <th style="border:1px solid #d4af37;padding:8px 10px;color:#d4af37;">Medio</th>
                    <th style="border:1px solid #d4af37;padding:8px 10px;color:#d4af37;">Alto</th>
                </tr>
                </thead>
                <tbody>
                <tr>
                    <td style="border:1px solid #445;padding:8px;font-weight:bold;background:#0a1628;color:#cbd5e1;">Alta</td>
                    <td style="border:1px solid #445;padding:8px;background:rgba(245,158,11,0.25);color:#f59e0b;font-weight:600;">Medio</td>
                    <td style="border:1px solid #445;padding:8px;background:rgba(220,38,38,0.25);color:#ef4444;font-weight:600;">Alto</td>
                    <td style="border:1px solid #445;padding:8px;background:rgba(220,38,38,0.45);color:#ff6b6b;font-weight:bold;">Critico</td>
                </tr>
                <tr>
                    <td style="border:1px solid #445;padding:8px;font-weight:bold;background:#0a1628;color:#cbd5e1;">Media</td>
                    <td style="border:1px solid #445;padding:8px;background:rgba(34,197,94,0.2);color:#22c55e;font-weight:600;">Baixo</td>
                    <td style="border:1px solid #445;padding:8px;background:rgba(245,158,11,0.25);color:#f59e0b;font-weight:600;">Medio</td>
                    <td style="border:1px solid #445;padding:8px;background:rgba(220,38,38,0.25);color:#ef4444;font-weight:600;">Alto</td>
                </tr>
                <tr>
                    <td style="border:1px solid #445;padding:8px;font-weight:bold;background:#0a1628;color:#cbd5e1;">Baixa</td>
                    <td style="border:1px solid #445;padding:8px;background:rgba(34,197,94,0.2);color:#22c55e;font-weight:600;">Baixo</td>
                    <td style="border:1px solid #445;padding:8px;background:rgba(34,197,94,0.2);color:#22c55e;font-weight:600;">Baixo</td>
                    <td style="border:1px solid #445;padding:8px;background:rgba(245,158,11,0.25);color:#f59e0b;font-weight:600;">Medio</td>
                </tr>
                </tbody>
                </table>
                </div>
                """
                st.markdown(legenda_html, unsafe_allow_html=True)
                st.caption("A tabela acima e uma referencia visual. O mapa detalhado por categoria esta na analise da IA.")
            else:
                st.warning("Descreva o cenário de contratação.")

# ============================================================
# TAB 3 — CHECKLIST DE PROCESSO
# ============================================================
with tab_checklist:
    st.markdown("### 📋 Checklist de Processo")
    st.markdown("Selecione a modalidade do processo e confira os documentos necessários. Marque os que já possui e imprima a capa do processo.")

    modalidade_sel = st.selectbox(
        "Modalidade / Tipo de processo",
        list(CHECKLIST_MODALIDADES.keys()),
        key="checklist_modalidade",
    )

    # Expander explicativo da modalidade
    info = INFO_MODALIDADES.get(modalidade_sel, {})
    if info:
        with st.expander(f"ℹ️ O que é: {modalidade_sel}", expanded=False):
            st.markdown(f"**📖 Descrição**\n\n{info['descricao']}")
            st.markdown(f"**📌 Quando se aplica**\n\n{info['quando_aplicar']}")
            st.markdown(f"**💰 Limites financeiros**\n\n{info['limites']}")
            st.markdown(f"**📋 Como montar o processo (Lei 14.133/2021)**\n\n{info['como_montar']}")

    itens_checklist = CHECKLIST_MODALIDADES[modalidade_sel]
    st.markdown(f"**{len(itens_checklist)} documentos necessários:**")
    st.markdown("---")

    status_checks = {}
    for item_ck in itens_checklist:
        _key = f"ck_{hashlib.md5((modalidade_sel + item_ck).encode()).hexdigest()[:10]}"
        status_checks[item_ck] = st.checkbox(item_ck, key=_key)

    # Resumo
    marcados = sum(1 for v in status_checks.values() if v)
    total = len(itens_checklist)
    st.markdown("---")
    if marcados == total:
        st.success(f"✅ Processo completo! {marcados}/{total} documentos presentes.")
    elif marcados > 0:
        st.warning(f"⚠️ {marcados}/{total} documentos presentes. Faltam {total - marcados}.")
    else:
        st.info(f"📋 {total} documentos necessários. Marque os que já possui.")

    # Botão imprimir PDF
    st.markdown("")
    if st.button("🖨️ Imprimir Checklist (PDF)", use_container_width=True, key="btn_print_checklist"):
        pdf_bytes = gerar_checklist_pdf(modalidade_sel, itens_checklist, status_checks)
        st.download_button(
            "⬇️ Baixar PDF",
            pdf_bytes,
            f"checklist_{modalidade_sel.replace(' ', '_').lower()}.pdf",
            "application/pdf",
            key="dl_checklist_pdf",
        )

# ============================================================
# TAB 4 — CÁLCULO IPCA
# ============================================================

@st.cache_data(ttl=86400, show_spinner=False)
def _buscar_ipca_bcb():
    """Busca série histórica completa do IPCA (série 433) no BCB."""
    url = "https://api.bcb.gov.br/dados/serie/bcdata.sgs.433/dados?formato=json"
    resp = requests.get(url, timeout=30)
    resp.raise_for_status()
    dados = resp.json()
    registros = []
    for r in dados:
        partes = r["data"].split("/")
        dt = datetime(int(partes[2]), int(partes[1]), int(partes[0]))
        registros.append({"data": dt, "valor": float(r["valor"])})
    return registros


def _calcular_ipca_acumulado(ipca_dados, data_inicio, data_fim):
    """Calcula o fator acumulado do IPCA entre duas datas (mês/ano)."""
    inicio_mes = datetime(data_inicio.year, data_inicio.month, 1)
    fim_mes = datetime(data_fim.year, data_fim.month, 1)
    fator = 1.0
    meses_usados = []
    for r in ipca_dados:
        if r["data"] >= inicio_mes and r["data"] <= fim_mes:
            fator *= (1 + r["valor"] / 100)
            meses_usados.append(r)
    percentual = (fator - 1) * 100
    return fator, percentual, meses_usados


def _gerar_pdf_ipca(itens_resultado, meses_detalhes, data_calc):
    """Gera relatório PDF detalhado do cálculo IPCA."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=25)

    # === CABEÇALHO ===
    pdf.set_fill_color(0, 26, 77)
    pdf.rect(0, 0, 210, 50, "F")
    pdf.set_draw_color(212, 175, 55)
    pdf.set_line_width(1.2)
    pdf.line(0, 50, 210, 50)

    pdf.set_y(8)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(212, 175, 55)
    pdf.cell(0, 5, "MARINHA DO BRASIL", ln=True, align="C")
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(200, 200, 220)
    pdf.cell(0, 4, "Centro de Operacoes do Abastecimento", ln=True, align="C")
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 9, _sanitize_for_pdf("RELATORIO DE CORRECAO PELO IPCA"), ln=True, align="C")
    pdf.ln(1)

    # === FAIXA DOURADA ===
    pdf.set_fill_color(212, 175, 55)
    pdf.rect(0, 52, 210, 14, "F")
    pdf.set_y(54)
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(0, 26, 77)
    pdf.cell(0, 9, _sanitize_for_pdf("Indice Nacional de Precos ao Consumidor Amplo"), ln=True, align="C")
    pdf.ln(6)

    # === METADADOS ===
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 5, f"Data/Hora do Calculo: {data_calc.strftime('%d/%m/%Y %H:%M:%S')}", ln=True)
    pdf.cell(0, 5, "API: Banco Central do Brasil - SGS Serie 433 (IPCA mensal)", ln=True)
    pdf.cell(0, 5, "URL: https://api.bcb.gov.br/dados/serie/bcdata.sgs.433/dados?formato=json", ln=True)
    pdf.cell(0, 5, "Fonte: IBGE - Instituto Brasileiro de Geografia e Estatistica", ln=True)
    pdf.ln(4)

    # === SEPARADOR ===
    pdf.set_draw_color(212, 175, 55)
    pdf.set_line_width(0.5)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(6)

    # === TABELA DE RESULTADOS ===
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(0, 26, 77)
    pdf.cell(0, 8, "ITENS CORRIGIDOS", ln=True, align="C")
    pdf.ln(3)

    # Cabeçalho tabela
    col_w = [12, 48, 28, 28, 28, 25, 25]
    headers = ["#", "Descricao", "Valor Orig.", "Data Orig.", "IPCA Acum.", "Fator", "Valor Corr."]
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_fill_color(0, 26, 77)
    pdf.set_text_color(255, 255, 255)
    for i, h in enumerate(headers):
        pdf.cell(col_w[i], 7, h, border=1, fill=True, align="C")
    pdf.ln()

    # Dados
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(30, 30, 30)
    total_original = 0.0
    total_corrigido = 0.0
    for idx, item in enumerate(itens_resultado, 1):
        total_original += item["valor_original"]
        total_corrigido += item["valor_corrigido"]
        fill = idx % 2 == 0
        if fill:
            pdf.set_fill_color(240, 245, 255)
        row = [
            str(idx),
            _sanitize_for_pdf(item["descricao"][:30]),
            f"R$ {item['valor_original']:,.2f}",
            item["data_original"].strftime("%m/%Y"),
            f"{item['percentual']:.4f}%",
            f"{item['fator']:.6f}",
            f"R$ {item['valor_corrigido']:,.2f}",
        ]
        for i, val in enumerate(row):
            pdf.cell(col_w[i], 6, val, border=1, fill=fill, align="C" if i != 1 else "L")
        pdf.ln()

    # Totais
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_fill_color(212, 175, 55)
    pdf.set_text_color(0, 26, 77)
    pdf.cell(sum(col_w[:2]), 7, "TOTAL", border=1, fill=True, align="C")
    pdf.cell(col_w[2], 7, f"R$ {total_original:,.2f}", border=1, fill=True, align="C")
    pdf.cell(sum(col_w[3:6]), 7, "", border=1, fill=True)
    pdf.cell(col_w[6], 7, f"R$ {total_corrigido:,.2f}", border=1, fill=True, align="C")
    pdf.ln()

    diferenca = total_corrigido - total_original
    pdf.ln(3)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(0, 26, 77)
    pdf.cell(0, 6, f"Diferenca total: R$ {diferenca:,.2f}", ln=True, align="R")
    pdf.ln(6)

    # === DETALHAMENTO MENSAL ===
    if meses_detalhes:
        pdf.set_draw_color(212, 175, 55)
        pdf.set_line_width(0.5)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())
        pdf.ln(4)

        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(0, 26, 77)
        pdf.cell(0, 7, "DETALHAMENTO MENSAL DO IPCA", ln=True, align="C")
        pdf.ln(3)

        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(0, 26, 77)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(40, 6, "Mes/Ano", border=1, fill=True, align="C")
        pdf.cell(40, 6, "IPCA Mensal (%)", border=1, fill=True, align="C")
        pdf.cell(50, 6, "Fator Acumulado", border=1, fill=True, align="C")
        pdf.ln()

        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(30, 30, 30)
        fator_acum = 1.0
        for i, m in enumerate(meses_detalhes):
            fator_acum *= (1 + m["valor"] / 100)
            fill = i % 2 == 0
            if fill:
                pdf.set_fill_color(240, 245, 255)
            pdf.cell(40, 5, m["data"].strftime("%m/%Y"), border=1, fill=fill, align="C")
            pdf.cell(40, 5, f"{m['valor']:.2f}%", border=1, fill=fill, align="C")
            pdf.cell(50, 5, f"{fator_acum:.6f}", border=1, fill=fill, align="C")
            pdf.ln()

    # === RODAPÉ ===
    pdf.ln(8)
    pdf.set_draw_color(212, 175, 55)
    pdf.set_line_width(0.3)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(3)
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 4, "Documento gerado automaticamente pelo sistema O Babilaca (IA)", ln=True, align="C")
    pdf.cell(0, 4, "Os valores sao meramente indicativos. Confirme sempre nas fontes oficiais.", ln=True, align="C")

    return bytes(pdf.output())


with tab_ipca:
    st.markdown("### 📊 Correção de Valores pelo IPCA")
    st.markdown(
        "Informe os itens com seus valores originais e datas de referência. "
        "O sistema buscará o IPCA acumulado no **Banco Central** e calculará o valor corrigido até o mês mais recente disponível."
    )

    # Inicializar itens na session_state
    if "ipca_itens" not in st.session_state:
        st.session_state["ipca_itens"] = [{"descricao": "", "valor": 0.0, "mes": 1, "ano": 2024}]

    # --- Formulário de entrada ---
    st.markdown("#### 📝 Itens para Correção")

    itens_ipca = st.session_state["ipca_itens"]

    for i, item in enumerate(itens_ipca):
        cols = st.columns([4, 2, 1, 1, 0.5])
        with cols[0]:
            itens_ipca[i]["descricao"] = st.text_input(
                "Descrição", value=item["descricao"], key=f"ipca_desc_{i}",
                placeholder="Ex: Material de expediente"
            )
        with cols[1]:
            itens_ipca[i]["valor"] = st.number_input(
                "Valor Original (R$)", value=item["valor"], min_value=0.0,
                format="%.2f", key=f"ipca_val_{i}", step=0.01
            )
        with cols[2]:
            itens_ipca[i]["mes"] = st.selectbox(
                "Mês", list(range(1, 13)), index=item["mes"] - 1, key=f"ipca_mes_{i}",
                format_func=lambda x: f"{x:02d}"
            )
        with cols[3]:
            itens_ipca[i]["ano"] = st.number_input(
                "Ano", value=item["ano"], min_value=1995, max_value=datetime.now().year,
                step=1, key=f"ipca_ano_{i}"
            )
        with cols[4]:
            st.markdown("<br>", unsafe_allow_html=True)
            if len(itens_ipca) > 1:
                if st.button("🗑️", key=f"ipca_rm_{i}", help="Remover item"):
                    itens_ipca.pop(i)
                    st.rerun()

    col_add, col_clear = st.columns([1, 1])
    with col_add:
        if st.button("➕ Adicionar Item", key="ipca_add_item", use_container_width=True):
            itens_ipca.append({"descricao": "", "valor": 0.0, "mes": 1, "ano": 2024})
            st.rerun()
    with col_clear:
        if st.button("🧹 Limpar Tudo", key="ipca_clear_all", use_container_width=True):
            st.session_state["ipca_itens"] = [{"descricao": "", "valor": 0.0, "mes": 1, "ano": 2024}]
            st.session_state.pop("ipca_resultado", None)
            st.rerun()

    st.markdown("---")

    # --- Botão calcular ---
    if st.button("🔢 Calcular Correção IPCA", type="primary", use_container_width=True, key="ipca_calc_btn"):
        itens_validos = [it for it in itens_ipca if it["descricao"].strip() and it["valor"] > 0]
        if not itens_validos:
            st.error("Preencha pelo menos um item com descrição e valor maior que zero.")
        else:
            with st.spinner("Buscando dados do IPCA no Banco Central..."):
                try:
                    ipca_dados = _buscar_ipca_bcb()
                    if not ipca_dados:
                        st.error("Não foi possível obter dados do IPCA. Tente novamente.")
                    else:
                        ultimo_mes = ipca_dados[-1]["data"]
                        data_calc = datetime.now()
                        resultados = []
                        todos_meses = []
                        for it in itens_validos:
                            dt_inicio = datetime(it["ano"], it["mes"], 1)
                            if dt_inicio > ultimo_mes:
                                st.warning(
                                    f"Item '{it['descricao']}': data {it['mes']:02d}/{it['ano']} "
                                    f"posterior ao último IPCA disponível ({ultimo_mes.strftime('%m/%Y')}). Ignorado."
                                )
                                continue
                            fator, percentual, meses = _calcular_ipca_acumulado(ipca_dados, dt_inicio, ultimo_mes)
                            valor_corrigido = it["valor"] * fator
                            resultados.append({
                                "descricao": it["descricao"],
                                "valor_original": it["valor"],
                                "data_original": dt_inicio,
                                "fator": fator,
                                "percentual": percentual,
                                "valor_corrigido": valor_corrigido,
                            })
                            if not todos_meses:
                                todos_meses = meses

                        if resultados:
                            st.session_state["ipca_resultado"] = {
                                "itens": resultados,
                                "meses": todos_meses,
                                "data_calc": data_calc,
                                "ultimo_ipca": ultimo_mes,
                            }
                except Exception as e:
                    st.error(f"Erro ao consultar API do BCB: {e}")

    # --- Exibição dos resultados ---
    if "ipca_resultado" in st.session_state:
        res = st.session_state["ipca_resultado"]
        itens_res = res["itens"]
        data_calc = res["data_calc"]
        ultimo_ipca = res["ultimo_ipca"]

        st.markdown("---")
        st.markdown("#### 📈 Resultado da Correção")
        st.caption(
            f"IPCA acumulado até **{ultimo_ipca.strftime('%m/%Y')}** · "
            f"Calculado em {data_calc.strftime('%d/%m/%Y às %H:%M:%S')} · "
            f"Fonte: BCB/IBGE (Série 433)"
        )

        # Tabela de resultados
        df_res = pd.DataFrame([
            {
                "Descrição": it["descricao"],
                "Valor Original": f"R$ {it['valor_original']:,.2f}",
                "Data Ref.": it["data_original"].strftime("%m/%Y"),
                "IPCA Acum. (%)": f"{it['percentual']:.4f}%",
                "Fator": f"{it['fator']:.6f}",
                "Valor Corrigido": f"R$ {it['valor_corrigido']:,.2f}",
            }
            for it in itens_res
        ])
        st.dataframe(df_res, use_container_width=True, hide_index=True)

        # Totais
        total_orig = sum(it["valor_original"] for it in itens_res)
        total_corr = sum(it["valor_corrigido"] for it in itens_res)
        diff = total_corr - total_orig

        c1, c2, c3 = st.columns(3)
        c1.metric("Total Original", f"R$ {total_orig:,.2f}")
        c2.metric("Total Corrigido", f"R$ {total_corr:,.2f}")
        c3.metric("Diferença", f"R$ {diff:,.2f}", delta=f"{(diff/total_orig*100) if total_orig else 0:.2f}%")

        # PDF
        st.markdown("---")
        pdf_ipca = _gerar_pdf_ipca(itens_res, res["meses"], data_calc)
        st.download_button(
            "📥 Baixar Relatório Detalhado (PDF)",
            pdf_ipca,
            f"relatorio_ipca_{data_calc.strftime('%Y%m%d_%H%M%S')}.pdf",
            "application/pdf",
            use_container_width=True,
            key="dl_ipca_pdf",
        )
